#!/usr/bin/env python3
"""
Story Bible MCP Server — shared multi-agent fiction authoring store.

Purpose:
    One SQLite-backed store that multiple AI clients (Claude, ChatGPT, anything
    MCP) read and write over HTTP MCP. Designed for an author/editor split:
    author-role keys have full write access; editor-role keys can read, comment,
    and file edit PROPOSALS but can never overwrite content directly.

Data model (schema.sql):
    projects → entities (arc/narrative/character/faction/lore/event/research/note)
             → chapters (prose, ordered, statused)
    revisions — immutable history; every content write bumps rev
    links     — typed edges between any two records (arc→character, ...)
    comments  — threaded, quote-anchored, open/resolved
    proposals — track-changes: editor proposes against a base rev; author
                accepts (creates new revision, stale-base guarded) or rejects

Auth:
    STORYBIBLE_KEYS env var: comma-separated "name:role:key" triples, e.g.
        STORYBIBLE_KEYS="joe:author:sk-abc,luna:author:sk-def,chatgpt:editor:sk-ghi"
    Clients send the key via X-API-Key header (or Authorization: Bearer).
    Every write is attributed to the key's name.

Inputs/outputs: JSON over MCP streamable-HTTP (stateless). GET /healthz is open.
Side effects: creates/updates STORYBIBLE_DB (default ~/.story-bible/story.db).
Failure behavior: unknown/missing key → 401 before any tool runs; tool-level
    errors return MCP errors with a human-readable message; the DB is opened
    per-request (WAL) so a crashed request never wedges the store.

Usage:
    STORYBIBLE_KEYS="..." python3 server.py            # serves on :8787
    claude mcp add --transport http story-bible https://host/mcp \
        --header "X-API-Key: sk-..."
"""

import contextvars
import json
import os
import re
import sqlite3
import sys
import threading
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path

from mcp.server.fastmcp import FastMCP
from mcp.server.transport_security import TransportSecuritySettings

DB_PATH = Path(os.environ.get("STORYBIBLE_DB", "~/.story-bible/story.db")).expanduser()
SCHEMA_PATH = Path(__file__).parent / "schema.sql"
# STORYBIBLE_PORT wins; PORT is honored for Railway/Heroku-style platforms that inject it.
PORT = int(os.environ.get("STORYBIBLE_PORT") or os.environ.get("PORT") or "8787")

EXPORT_DIR = DB_PATH.parent / "exports"

# On-volume snapshot settings. BACKUP_HOURS <= 0 disables the timer (tools still work).
BACKUP_DIR = DB_PATH.parent / "backups"
BACKUP_HOURS = float(os.environ.get("STORYBIBLE_BACKUP_HOURS", "24"))
BACKUP_KEEP = int(os.environ.get("STORYBIBLE_BACKUP_KEEP", "14"))

ENTITY_KINDS = {"arc", "narrative", "character", "faction", "lore", "event", "research", "note",
                "location", "item", "theme", "thread", "style",
                "voice_profile", "narrative_profile", "rubric_profile", "scene_intent"}
CHAPTER_STATUSES = {"draft", "revised", "final"}
SCENE_STATUSES = {"outline", "draft", "revised", "final"}
TARGET_TABLES = {"entity": "entities", "chapter": "chapters", "scene": "scenes"}
META_TARGETS = {"project", "entity", "chapter", "scene"}
ROLES = {"author", "editor", "owner"}
NARRATIVE_MODES = {"memory", "direct", "documentary", "institutional", "transcript",
                   "artifact", "omniscient", "mixed", ""}
VISIBILITIES = {"public", "private", "silhouette"}
VERDICTS = {"verified", "false", "disputed", "unverifiable", "outdated"}
CLAIM_DOMAINS = {"technical", "scientific", "political", "historical", "geographic",
                 "medical", "legal", "economic", "cultural", "other"}
CLAIM_CLASSES = {"fact", "inference", "projection", "deliberate_fiction"}
ANALYSIS_TYPES = {"scene_check", "chapter_gate", "part_audit", "global_arc_audit",
                  "fact_check", "second_opinion", "cold_read"}
ADVISORY_TYPES = {"second_opinion", "cold_read"}
SEVERITIES = {"blocking", "major", "minor", "watch"}
FINDING_STATUSES = {"open", "accepted", "resolved", "intentional", "deferred", "incorrect"}
OWNER_FINDING_STATUSES = {"intentional", "incorrect"}

# caller identity for the current request: {"name": ..., "role": "author"|"editor"}
CALLER: contextvars.ContextVar[dict] = contextvars.ContextVar("caller")

# DNS-rebinding protection is for unauthenticated localhost servers; this one is
# API-key-gated and internet-facing (Railway edge sets Host to the public domain,
# which the SDK's default localhost allowlist 421s). Disabled unless
# STORYBIBLE_ALLOWED_HOSTS provides an explicit allowlist.
_allowed_hosts = [h.strip() for h in os.environ.get("STORYBIBLE_ALLOWED_HOSTS", "").split(",") if h.strip()]
_transport_security = (
    TransportSecuritySettings(enable_dns_rebinding_protection=True, allowed_hosts=_allowed_hosts)
    if _allowed_hosts
    else TransportSecuritySettings(enable_dns_rebinding_protection=False)
)

mcp = FastMCP("story-bible", stateless_http=True, json_response=True,
              transport_security=_transport_security)


# ---------------------------------------------------------------- helpers

def _load_keys() -> dict:
    """Parse STORYBIBLE_KEYS into {key: {"name":..., "role":...}}."""
    raw = os.environ.get("STORYBIBLE_KEYS", "")
    keys = {}
    for triple in filter(None, (t.strip() for t in raw.split(","))):
        try:
            name, role, key = triple.split(":", 2)
        except ValueError:
            print(f"[story-bible] bad key entry (want name:role:key): {triple!r}", file=sys.stderr)
            continue
        if role not in ROLES:
            print(f"[story-bible] bad role {role!r} for {name!r}", file=sys.stderr)
            continue
        keys[key] = {"name": name, "role": role}
    return keys


def _db() -> sqlite3.Connection:
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def _init_db():
    with _db() as conn:
        conn.executescript(SCHEMA_PATH.read_text())
        # One-time FTS backfill for databases created before the fts table existed.
        if conn.execute("SELECT COUNT(*) n FROM fts").fetchone()["n"] == 0:
            conn.execute(
                "INSERT INTO fts(target_type, target_id, project_id, name, body) "
                "SELECT 'entity', id, project_id, name || ' ' || summary, content_md "
                "FROM entities WHERE deleted=0")
            conn.execute(
                "INSERT INTO fts(target_type, target_id, project_id, name, body) "
                "SELECT 'chapter', id, project_id, title, content_md "
                "FROM chapters WHERE deleted=0")
            conn.execute(
                "INSERT INTO fts(target_type, target_id, project_id, name, body) "
                "SELECT 'scene', id, project_id, title || ' ' || synopsis, content_md "
                "FROM scenes WHERE deleted=0")
        # Guarded column additions for databases created before these existed.
        for stmt in ("ALTER TABLE links ADD COLUMN attrs TEXT DEFAULT ''",
                     "ALTER TABLE projects ADD COLUMN deleted INTEGER NOT NULL DEFAULT 0",
                     "ALTER TABLE chapters ADD COLUMN part_id TEXT",
                     "ALTER TABLE entities ADD COLUMN visibility TEXT NOT NULL DEFAULT 'public'",
                     "ALTER TABLE scenes ADD COLUMN narrative_mode TEXT DEFAULT ''",
                     "ALTER TABLE scenes ADD COLUMN story_time_start TEXT DEFAULT ''",
                     "ALTER TABLE scenes ADD COLUMN story_time_end TEXT DEFAULT ''",
                     "ALTER TABLE oauth_codes ADD COLUMN client_id TEXT DEFAULT ''"):
            try:
                conn.execute(stmt)
            except sqlite3.OperationalError as e:
                if "duplicate column" not in str(e).lower():
                    raise


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _id() -> str:
    return str(uuid.uuid4())


def _caller() -> dict:
    return CALLER.get({"name": "unknown", "role": "editor"})


def _require_author():
    c = _caller()
    if c["role"] not in ("author", "owner"):
        raise PermissionError(
            f"'{c['name']}' has editor role: read, comment, and propose only. "
            "Use proposal_create to suggest this change."
        )


def _require_owner():
    c = _caller()
    if c["role"] != "owner":
        raise PermissionError(f"'{c['name']}' is not the owner — this action is owner-only.")


def _log_lock_event(conn, target_type: str, target_id: str, action: str, detail: str = ""):
    conn.execute(
        "INSERT INTO lock_events (id, target_type, target_id, action, attempted_by, detail, created_at) "
        "VALUES (?,?,?,?,?,?,?)",
        (_id(), target_type, target_id, action, _caller()["name"], detail, _now()))


def _require_unlocked(conn, target_type: str, target_id: str, action: str = "write"):
    """Locked targets reject every mutation; the attempt itself is recorded (spec 3.12/F)."""
    row = conn.execute("SELECT kind, reason, locked_by FROM locks WHERE target_type=? AND target_id=?",
                       (target_type, target_id)).fetchone()
    if row is not None:
        # Log on an independent connection: the caller's transaction rolls back with
        # the raised error, and the audit row must survive the refusal.
        log_conn = _db()
        try:
            with log_conn:
                log_conn.execute(
                    "INSERT INTO lock_events (id, target_type, target_id, action, attempted_by, "
                    "detail, created_at) VALUES (?,?,?,?,?,?,?)",
                    (_id(), target_type, target_id, "blocked_write", _caller()["name"],
                     action, _now()))
        finally:
            log_conn.close()
        raise PermissionError(
            f"{target_type} {target_id} is locked ({row['kind']}"
            + (f": {row['reason']}" if row["reason"] else "")
            + f", by {row['locked_by']}). Only an owner decision changes locked records.")


def _row(r: sqlite3.Row | None) -> dict | None:
    return dict(r) if r is not None else None


def _rows(rs) -> list[dict]:
    return [dict(r) for r in rs]


def _get_target(conn, target_type: str, target_id: str) -> sqlite3.Row:
    if target_type not in TARGET_TABLES:
        raise ValueError(f"target_type must be one of {sorted(TARGET_TABLES)}")
    table = TARGET_TABLES[target_type]
    row = conn.execute(f"SELECT * FROM {table} WHERE id=? AND deleted=0", (target_id,)).fetchone()
    if row is None:
        raise ValueError(f"{target_type} {target_id} not found")
    return row


def _write_revision(conn, target_type: str, target_id: str, content_md: str,
                    note: str, created_by: str) -> int:
    """Bump rev on the target row and record an immutable revision. Returns new rev."""
    table = TARGET_TABLES[target_type]
    cur = conn.execute(f"SELECT rev FROM {table} WHERE id=?", (target_id,)).fetchone()
    new_rev = cur["rev"] + 1
    conn.execute(
        f"UPDATE {table} SET content_md=?, rev=?, updated_at=? WHERE id=?",
        (content_md, new_rev, _now(), target_id),
    )
    conn.execute(
        "INSERT INTO revisions (id, target_type, target_id, rev, content_md, note, created_by, created_at) "
        "VALUES (?,?,?,?,?,?,?,?)",
        (_id(), target_type, target_id, new_rev, content_md, note, created_by, _now()),
    )
    _refresh_mentions(conn, target_type, target_id)
    return new_rev


def _record_initial_revision(conn, target_type: str, target_id: str, content_md: str, created_by: str):
    conn.execute(
        "INSERT INTO revisions (id, target_type, target_id, rev, content_md, note, created_by, created_at) "
        "VALUES (?,?,?,1,?,'initial',?,?)",
        (_id(), target_type, target_id, content_md, created_by, _now()),
    )
    _refresh_mentions(conn, target_type, target_id)


# ---------------------------------------------------------------- mention index

def _entity_name_map(conn, project_id: str) -> dict[str, list[str]]:
    """entity_id -> lowercase [name, *aliases] (min length 3 to avoid noise words)."""
    out: dict[str, list[str]] = {}
    for r in conn.execute("SELECT id, name FROM entities WHERE project_id=? AND deleted=0",
                          (project_id,)):
        names = [r["name"]]
        al = conn.execute(
            "SELECT value FROM node_meta WHERE target_type='entity' AND target_id=? AND key='aliases'",
            (r["id"],)).fetchone()
        if al is not None:
            names += [a.strip() for a in al["value"].split(",")]
        cleaned = [n.lower() for n in names if len(n.strip()) >= 3]
        if cleaned:
            out[r["id"]] = cleaned
    return out


def _refresh_mentions(conn, target_type: str, target_id: str):
    """Rescan one node's current content for entity name/alias occurrences (word-bounded)."""
    row = conn.execute(
        f"SELECT project_id, content_md FROM {TARGET_TABLES[target_type]} WHERE id=?",
        (target_id,)).fetchone()
    if row is None:
        return
    text = (row["content_md"] or "").lower()
    conn.execute("DELETE FROM mentions WHERE target_type=? AND target_id=?",
                 (target_type, target_id))
    if not text:
        return
    for eid, names in _entity_name_map(conn, row["project_id"]).items():
        if target_type == "entity" and eid == target_id:
            continue  # a record doesn't mention itself
        n = sum(len(re.findall(rf"(?<!\w){re.escape(nm)}(?!\w)", text)) for nm in names)
        if n:
            conn.execute(
                "INSERT INTO mentions (project_id, entity_id, target_type, target_id, count, updated_at) "
                "VALUES (?,?,?,?,?,?)",
                (row["project_id"], eid, target_type, target_id, n, _now()))


# ---------------------------------------------------------------- projects

@mcp.tool()
def project_create(name: str, description: str = "") -> dict:
    """Create a project (a book or series). Author role required."""
    _require_author()
    with _db() as conn:
        pid = _id()
        conn.execute(
            "INSERT INTO projects (id, name, description, created_at, updated_at) VALUES (?,?,?,?,?)",
            (pid, name, description, _now(), _now()),
        )
        return _row(conn.execute("SELECT * FROM projects WHERE id=?", (pid,)).fetchone())


@mcp.tool()
def project_list() -> list[dict]:
    """List all projects."""
    with _db() as conn:
        return _rows(conn.execute("SELECT * FROM projects WHERE deleted=0 ORDER BY created_at"))


@mcp.tool()
def project_delete(project_id: str) -> dict:
    """Soft-delete a project — it disappears from listings but nothing is destroyed
    (entities, chapters, revisions all retained). Author role required."""
    _require_author()
    with _db() as conn:
        if conn.execute("SELECT 1 FROM projects WHERE id=? AND deleted=0",
                        (project_id,)).fetchone() is None:
            raise ValueError(f"project {project_id} not found")
        conn.execute("UPDATE projects SET deleted=1, updated_at=? WHERE id=?",
                     (_now(), project_id))
        return {"deleted": project_id}


@mcp.tool()
def project_get(project_id: str) -> dict:
    """Get a project with entity/chapter counts by kind."""
    with _db() as conn:
        proj = _row(conn.execute("SELECT * FROM projects WHERE id=? AND deleted=0",
                                 (project_id,)).fetchone())
        if proj is None:
            raise ValueError(f"project {project_id} not found")
        proj["entity_counts"] = {
            r["kind"]: r["n"]
            for r in conn.execute(
                "SELECT kind, COUNT(*) n FROM entities WHERE project_id=? AND deleted=0 GROUP BY kind",
                (project_id,),
            )
        }
        proj["chapter_count"] = conn.execute(
            "SELECT COUNT(*) n FROM chapters WHERE project_id=? AND deleted=0", (project_id,)
        ).fetchone()["n"]
        return proj


@mcp.tool()
def project_update(project_id: str, name: str | None = None, description: str | None = None) -> dict:
    """Update project name/description. Author role required."""
    _require_author()
    with _db() as conn:
        if conn.execute("SELECT 1 FROM projects WHERE id=?", (project_id,)).fetchone() is None:
            raise ValueError(f"project {project_id} not found")
        if name is not None:
            conn.execute("UPDATE projects SET name=?, updated_at=? WHERE id=?", (name, _now(), project_id))
        if description is not None:
            conn.execute("UPDATE projects SET description=?, updated_at=? WHERE id=?",
                         (description, _now(), project_id))
        return _row(conn.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone())


# ---------------------------------------------------------------- entities

@mcp.tool()
def entity_create(project_id: str, kind: str, name: str, summary: str = "",
                  content_md: str = "", sort_order: int = 0) -> dict:
    """Create a structural entity. kind: arc | narrative | character | faction |
    lore | event | research | note. Author role required."""
    _require_author()
    if kind not in ENTITY_KINDS:
        raise ValueError(f"kind must be one of {sorted(ENTITY_KINDS)}")
    who = _caller()["name"]
    with _db() as conn:
        if conn.execute("SELECT 1 FROM projects WHERE id=?", (project_id,)).fetchone() is None:
            raise ValueError(f"project {project_id} not found")
        eid = _id()
        conn.execute(
            "INSERT INTO entities (id, project_id, kind, name, summary, content_md, sort_order, "
            "created_by, created_at, updated_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (eid, project_id, kind, name, summary, content_md, sort_order, who, _now(), _now()),
        )
        _record_initial_revision(conn, "entity", eid, content_md, who)
        return _row(conn.execute("SELECT * FROM entities WHERE id=?", (eid,)).fetchone())


@mcp.tool()
def entity_get(entity_id: str) -> dict:
    """Get an entity with its links and open comment count."""
    with _db() as conn:
        ent = _row(_get_target(conn, "entity", entity_id))
        ent["links"] = _rows(conn.execute(
            "SELECT * FROM links WHERE from_id=? OR to_id=?", (entity_id, entity_id)))
        ent["open_comments"] = conn.execute(
            "SELECT COUNT(*) n FROM comments WHERE target_type='entity' AND target_id=? AND status='open'",
            (entity_id,),
        ).fetchone()["n"]
        return ent


@mcp.tool()
def entity_list(project_id: str, kind: str | None = None) -> list[dict]:
    """List entities in a project, optionally filtered by kind. Content omitted; use entity_get."""
    with _db() as conn:
        q = ("SELECT id, kind, name, summary, rev, sort_order, created_by, updated_at "
             "FROM entities WHERE project_id=? AND deleted=0")
        args: list = [project_id]
        if kind is not None:
            if kind not in ENTITY_KINDS:
                raise ValueError(f"kind must be one of {sorted(ENTITY_KINDS)}")
            q += " AND kind=?"
            args.append(kind)
        return _rows(conn.execute(q + " ORDER BY kind, sort_order, name", args))


@mcp.tool()
def entity_update(entity_id: str, name: str | None = None, summary: str | None = None,
                  content_md: str | None = None, sort_order: int | None = None,
                  revision_note: str = "") -> dict:
    """Update an entity (partial patch). A content_md change records a new revision.
    Author role required."""
    _require_author()
    who = _caller()["name"]
    with _db() as conn:
        _get_target(conn, "entity", entity_id)
        _require_unlocked(conn, "entity", entity_id, "entity_update")
        for field, val in (("name", name), ("summary", summary), ("sort_order", sort_order)):
            if val is not None:
                conn.execute(f"UPDATE entities SET {field}=?, updated_at=? WHERE id=?",
                             (val, _now(), entity_id))
        if content_md is not None:
            _write_revision(conn, "entity", entity_id, content_md, revision_note, who)
        return _row(conn.execute("SELECT * FROM entities WHERE id=?", (entity_id,)).fetchone())


@mcp.tool()
def entity_delete(entity_id: str) -> dict:
    """Soft-delete an entity (revisions and comments are retained). Author role required."""
    _require_author()
    with _db() as conn:
        _get_target(conn, "entity", entity_id)
        _require_unlocked(conn, "entity", entity_id, "entity_delete")
        conn.execute("UPDATE entities SET deleted=1, updated_at=? WHERE id=?", (_now(), entity_id))
        return {"deleted": entity_id}


# ---------------------------------------------------------------- links

@mcp.tool()
def link_create(project_id: str, from_id: str, to_id: str, rel_type: str, note: str = "",
                attrs: dict | None = None) -> dict:
    """Link two records with a typed, directional relationship (e.g. character 'member_of'
    faction, arc 'centers_on' character). attrs is an optional structured payload for
    richer semantics — suggested keys: role (participant|observer), strength (1-5),
    since / until (story dates), inverse (label for the reverse direction).
    Author role required."""
    _require_author()
    with _db() as conn:
        lid = _id()
        conn.execute(
            "INSERT INTO links (id, project_id, from_id, to_id, rel_type, note, attrs, "
            "created_by, created_at) VALUES (?,?,?,?,?,?,?,?,?)",
            (lid, project_id, from_id, to_id, rel_type, note,
             json.dumps(attrs) if attrs else "", _caller()["name"], _now()),
        )
        return _row(conn.execute("SELECT * FROM links WHERE id=?", (lid,)).fetchone())


@mcp.tool()
def link_list(project_id: str, record_id: str | None = None) -> list[dict]:
    """List links in a project, optionally only those touching record_id."""
    with _db() as conn:
        if record_id is None:
            return _rows(conn.execute("SELECT * FROM links WHERE project_id=?", (project_id,)))
        return _rows(conn.execute(
            "SELECT * FROM links WHERE project_id=? AND (from_id=? OR to_id=?)",
            (project_id, record_id, record_id)))


@mcp.tool()
def link_delete(link_id: str) -> dict:
    """Delete a link. Author role required."""
    _require_author()
    with _db() as conn:
        if conn.execute("SELECT 1 FROM links WHERE id=?", (link_id,)).fetchone() is None:
            raise ValueError(f"link {link_id} not found")
        conn.execute("DELETE FROM links WHERE id=?", (link_id,))
        return {"deleted": link_id}


# ---------------------------------------------------------------- chapters

@mcp.tool()
def chapter_create(project_id: str, title: str, content_md: str = "", sort_order: int = 0) -> dict:
    """Create a prose chapter. Author role required."""
    _require_author()
    who = _caller()["name"]
    with _db() as conn:
        if conn.execute("SELECT 1 FROM projects WHERE id=?", (project_id,)).fetchone() is None:
            raise ValueError(f"project {project_id} not found")
        cid = _id()
        conn.execute(
            "INSERT INTO chapters (id, project_id, title, content_md, sort_order, "
            "created_by, created_at, updated_at) VALUES (?,?,?,?,?,?,?,?)",
            (cid, project_id, title, content_md, sort_order, who, _now(), _now()),
        )
        _record_initial_revision(conn, "chapter", cid, content_md, who)
        return _row(conn.execute("SELECT * FROM chapters WHERE id=?", (cid,)).fetchone())


@mcp.tool()
def chapter_get(chapter_id: str) -> dict:
    """Get a chapter including current content, its scene list (if it uses scenes),
    open comments count, and pending proposals count."""
    with _db() as conn:
        ch = _row(_get_target(conn, "chapter", chapter_id))
        ch["scenes"] = _rows(conn.execute(
            "SELECT id, title, synopsis, status, pov_entity_id, sort_order, rev, "
            "length(content_md) AS content_chars "
            "FROM scenes WHERE chapter_id=? AND deleted=0 ORDER BY sort_order, created_at",
            (chapter_id,)))
        ch["open_comments"] = conn.execute(
            "SELECT COUNT(*) n FROM comments WHERE target_type='chapter' AND target_id=? AND status='open'",
            (chapter_id,)).fetchone()["n"]
        ch["pending_proposals"] = conn.execute(
            "SELECT COUNT(*) n FROM proposals WHERE target_type='chapter' AND target_id=? AND status='pending'",
            (chapter_id,)).fetchone()["n"]
        return ch


@mcp.tool()
def chapter_list(project_id: str) -> list[dict]:
    """List chapters in order (content omitted; use chapter_get)."""
    with _db() as conn:
        return _rows(conn.execute(
            "SELECT id, title, status, rev, sort_order, created_by, updated_at, length(content_md) AS content_chars "
            "FROM chapters WHERE project_id=? AND deleted=0 ORDER BY sort_order, created_at",
            (project_id,)))


@mcp.tool()
def chapter_update(chapter_id: str, title: str | None = None, content_md: str | None = None,
                   status: str | None = None, sort_order: int | None = None,
                   part_id: str | None = None, revision_note: str = "",
                   final_override: bool = False) -> dict:
    """Update a chapter (partial patch). content_md change records a new revision.
    Setting status='final' runs the deterministic final gate (voice lint, silhouette
    leaks, false facts); a blocked gate needs fixes, or an owner passing
    final_override=true (which logs a decision). Pass part_id="" to unassign.
    Author role required."""
    _require_author()
    who = _caller()["name"]
    with _db() as conn:
        _get_target(conn, "chapter", chapter_id)
        _require_unlocked(conn, "chapter", chapter_id, "chapter_update")
        if status is not None and status not in CHAPTER_STATUSES:
            raise ValueError(f"status must be one of {sorted(CHAPTER_STATUSES)}")
        if part_id:
            ch = conn.execute("SELECT project_id FROM chapters WHERE id=?", (chapter_id,)).fetchone()
            p = conn.execute("SELECT project_id FROM parts WHERE id=? AND deleted=0",
                             (part_id,)).fetchone()
            if p is None or p["project_id"] != ch["project_id"]:
                raise ValueError(f"part {part_id} not found in this project")
        if content_md is not None:
            _write_revision(conn, "chapter", chapter_id, content_md, revision_note, who)
        if status == "final":
            _gate_final(conn, "chapter", chapter_id, final_override)
        for field, val in (("title", title), ("status", status), ("sort_order", sort_order)):
            if val is not None:
                conn.execute(f"UPDATE chapters SET {field}=?, updated_at=? WHERE id=?",
                             (val, _now(), chapter_id))
        if part_id is not None:
            conn.execute("UPDATE chapters SET part_id=?, updated_at=? WHERE id=?",
                         (part_id or None, _now(), chapter_id))
        return _row(conn.execute("SELECT * FROM chapters WHERE id=?", (chapter_id,)).fetchone())


@mcp.tool()
def chapter_delete(chapter_id: str) -> dict:
    """Soft-delete a chapter (revisions and comments are retained). Author role required."""
    _require_author()
    with _db() as conn:
        _get_target(conn, "chapter", chapter_id)
        _require_unlocked(conn, "chapter", chapter_id, "chapter_delete")
        conn.execute("UPDATE chapters SET deleted=1, updated_at=? WHERE id=?", (_now(), chapter_id))
        return {"deleted": chapter_id}


# ---------------------------------------------------------------- scenes

def _check_pov(conn, pov_entity_id: str | None, project_id: str):
    if pov_entity_id:
        row = conn.execute("SELECT kind, project_id FROM entities WHERE id=? AND deleted=0",
                           (pov_entity_id,)).fetchone()
        if row is None:
            raise ValueError(f"pov_entity_id {pov_entity_id} not found")
        if row["project_id"] != project_id:
            raise ValueError("pov_entity_id belongs to a different project")
        if row["kind"] != "character":
            raise ValueError("pov_entity_id must be a character entity")


@mcp.tool()
def scene_create(chapter_id: str, title: str = "", synopsis: str = "", content_md: str = "",
                 status: str = "outline", pov_entity_id: str | None = None,
                 sort_order: int = 0) -> dict:
    """Create a scene inside a chapter — the atomic prose/metadata unit. synopsis is the
    index-card summary; pov_entity_id points at the POV character entity; status:
    outline | draft | revised | final. Author role required."""
    _require_author()
    if status not in SCENE_STATUSES:
        raise ValueError(f"status must be one of {sorted(SCENE_STATUSES)}")
    who = _caller()["name"]
    with _db() as conn:
        ch = _get_target(conn, "chapter", chapter_id)
        _check_pov(conn, pov_entity_id, ch["project_id"])
        sid = _id()
        conn.execute(
            "INSERT INTO scenes (id, project_id, chapter_id, title, synopsis, content_md, "
            "status, pov_entity_id, sort_order, created_by, created_at, updated_at) "
            "VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
            (sid, ch["project_id"], chapter_id, title, synopsis, content_md, status,
             pov_entity_id, sort_order, who, _now(), _now()),
        )
        _record_initial_revision(conn, "scene", sid, content_md, who)
        return _row(conn.execute("SELECT * FROM scenes WHERE id=?", (sid,)).fetchone())


@mcp.tool()
def scene_get(scene_id: str) -> dict:
    """Get a scene with full content, metadata, open comment count, and pending proposal count."""
    with _db() as conn:
        sc = _row(_get_target(conn, "scene", scene_id))
        sc["meta"] = {r["key"]: r["value"] for r in conn.execute(
            "SELECT key, value FROM node_meta WHERE target_type='scene' AND target_id=?",
            (scene_id,))}
        sc["open_comments"] = conn.execute(
            "SELECT COUNT(*) n FROM comments WHERE target_type='scene' AND target_id=? AND status='open'",
            (scene_id,)).fetchone()["n"]
        sc["pending_proposals"] = conn.execute(
            "SELECT COUNT(*) n FROM proposals WHERE target_type='scene' AND target_id=? AND status='pending'",
            (scene_id,)).fetchone()["n"]
        return sc


@mcp.tool()
def scene_list(project_id: str | None = None, chapter_id: str | None = None,
               status: str | None = None, pov_entity_id: str | None = None) -> list[dict]:
    """List scenes (content omitted; use scene_get) filtered by chapter or project, and
    optionally by status and/or POV entity. Ordered by chapter, then scene order."""
    if project_id is None and chapter_id is None:
        raise ValueError("provide project_id or chapter_id")
    with _db() as conn:
        q = ("SELECT s.id, s.chapter_id, c.title AS chapter_title, s.title, s.synopsis, "
             "s.status, s.pov_entity_id, s.sort_order, s.rev, s.created_by, s.updated_at, "
             "length(s.content_md) AS content_chars "
             "FROM scenes s JOIN chapters c ON c.id = s.chapter_id WHERE s.deleted=0")
        args: list = []
        if chapter_id is not None:
            q += " AND s.chapter_id=?"
            args.append(chapter_id)
        if project_id is not None:
            q += " AND s.project_id=?"
            args.append(project_id)
        if status is not None:
            if status not in SCENE_STATUSES:
                raise ValueError(f"status must be one of {sorted(SCENE_STATUSES)}")
            q += " AND s.status=?"
            args.append(status)
        if pov_entity_id is not None:
            q += " AND s.pov_entity_id=?"
            args.append(pov_entity_id)
        return _rows(conn.execute(
            q + " ORDER BY c.sort_order, c.created_at, s.sort_order, s.created_at", args))


@mcp.tool()
def scene_update(scene_id: str, title: str | None = None, synopsis: str | None = None,
                 content_md: str | None = None, status: str | None = None,
                 pov_entity_id: str | None = None, sort_order: int | None = None,
                 revision_note: str = "") -> dict:
    """Update a scene (partial patch). A content_md change records a new revision.
    Pass pov_entity_id="" to clear it. Author role required."""
    _require_author()
    who = _caller()["name"]
    with _db() as conn:
        sc = _get_target(conn, "scene", scene_id)
        _require_unlocked(conn, "scene", scene_id, "scene_update")
        if status is not None and status not in SCENE_STATUSES:
            raise ValueError(f"status must be one of {sorted(SCENE_STATUSES)}")
        if pov_entity_id:
            _check_pov(conn, pov_entity_id, sc["project_id"])
        if content_md is not None:
            _write_revision(conn, "scene", scene_id, content_md, revision_note, who)
        if status == "final":
            _gate_final(conn, "scene", scene_id, False)
        for field, val in (("title", title), ("synopsis", synopsis), ("status", status),
                           ("sort_order", sort_order)):
            if val is not None:
                conn.execute(f"UPDATE scenes SET {field}=?, updated_at=? WHERE id=?",
                             (val, _now(), scene_id))
        if pov_entity_id is not None:
            conn.execute("UPDATE scenes SET pov_entity_id=?, updated_at=? WHERE id=?",
                         (pov_entity_id or None, _now(), scene_id))
        return _row(conn.execute("SELECT * FROM scenes WHERE id=?", (scene_id,)).fetchone())


@mcp.tool()
def scene_move(scene_id: str, chapter_id: str, sort_order: int | None = None) -> dict:
    """Move a scene to another chapter (same project), optionally setting its position.
    Author role required."""
    _require_author()
    with _db() as conn:
        sc = _get_target(conn, "scene", scene_id)
        _require_unlocked(conn, "scene", scene_id, "scene_move")
        ch = _get_target(conn, "chapter", chapter_id)
        if ch["project_id"] != sc["project_id"]:
            raise ValueError("target chapter is in a different project")
        conn.execute("UPDATE scenes SET chapter_id=?, updated_at=? WHERE id=?",
                     (chapter_id, _now(), scene_id))
        if sort_order is not None:
            conn.execute("UPDATE scenes SET sort_order=? WHERE id=?", (sort_order, scene_id))
        return _row(conn.execute("SELECT * FROM scenes WHERE id=?", (scene_id,)).fetchone())


@mcp.tool()
def scene_delete(scene_id: str) -> dict:
    """Soft-delete a scene (revisions and comments are retained). Author role required."""
    _require_author()
    with _db() as conn:
        _get_target(conn, "scene", scene_id)
        _require_unlocked(conn, "scene", scene_id, "scene_delete")
        conn.execute("UPDATE scenes SET deleted=1, updated_at=? WHERE id=?", (_now(), scene_id))
        return {"deleted": scene_id}


# ---------------------------------------------------------------- node metadata

@mcp.tool()
def meta_set(target_type: str, target_id: str, key: str, value: str) -> dict:
    """Set a metadata key on a project/entity/chapter/scene (tags, aliases, story_date,
    target_words, tone — any string key/value). Empty value deletes the key.
    Author role required."""
    _require_author()
    if target_type not in META_TARGETS:
        raise ValueError(f"target_type must be one of {sorted(META_TARGETS)}")
    if not key.strip():
        raise ValueError("key must be non-empty")
    with _db() as conn:
        if target_type == "project":
            if conn.execute("SELECT 1 FROM projects WHERE id=?", (target_id,)).fetchone() is None:
                raise ValueError(f"project {target_id} not found")
        else:
            _get_target(conn, target_type, target_id)
            _require_unlocked(conn, target_type, target_id, f"meta_set:{key}")
        if value == "":
            conn.execute("DELETE FROM node_meta WHERE target_type=? AND target_id=? AND key=?",
                         (target_type, target_id, key))
            return {"target_type": target_type, "target_id": target_id, "key": key,
                    "deleted": True}
        conn.execute(
            "INSERT INTO node_meta (target_type, target_id, key, value, updated_by, updated_at) "
            "VALUES (?,?,?,?,?,?) ON CONFLICT(target_type, target_id, key) "
            "DO UPDATE SET value=excluded.value, updated_by=excluded.updated_by, "
            "updated_at=excluded.updated_at",
            (target_type, target_id, key, value, _caller()["name"], _now()),
        )
        return {"target_type": target_type, "target_id": target_id, "key": key, "value": value}


@mcp.tool()
def meta_get(target_type: str, target_id: str) -> dict:
    """Get all metadata key/values on a node. Any role."""
    if target_type not in META_TARGETS:
        raise ValueError(f"target_type must be one of {sorted(META_TARGETS)}")
    with _db() as conn:
        return {r["key"]: r["value"] for r in conn.execute(
            "SELECT key, value FROM node_meta WHERE target_type=? AND target_id=? ORDER BY key",
            (target_type, target_id))}


# ---------------------------------------------------------------- revisions

@mcp.tool()
def revision_list(target_type: str, target_id: str) -> list[dict]:
    """List all revisions of an entity or chapter, newest first (content omitted)."""
    with _db() as conn:
        _get_target(conn, target_type, target_id)
        return _rows(conn.execute(
            "SELECT id, rev, note, created_by, created_at, length(content_md) AS content_chars "
            "FROM revisions WHERE target_type=? AND target_id=? ORDER BY rev DESC",
            (target_type, target_id)))


@mcp.tool()
def revision_get(revision_id: str) -> dict:
    """Get a specific revision including its full content."""
    with _db() as conn:
        rev = _row(conn.execute("SELECT * FROM revisions WHERE id=?", (revision_id,)).fetchone())
        if rev is None:
            raise ValueError(f"revision {revision_id} not found")
        return rev


@mcp.tool()
def revision_restore(revision_id: str) -> dict:
    """Restore an old revision by copying it forward as a NEW revision (history is never
    rewritten). Author role required."""
    _require_author()
    who = _caller()["name"]
    with _db() as conn:
        old = conn.execute("SELECT * FROM revisions WHERE id=?", (revision_id,)).fetchone()
        if old is None:
            raise ValueError(f"revision {revision_id} not found")
        _get_target(conn, old["target_type"], old["target_id"])
        _require_unlocked(conn, old["target_type"], old["target_id"], "revision_restore")
        new_rev = _write_revision(conn, old["target_type"], old["target_id"], old["content_md"],
                                  f"restored from rev {old['rev']}", who)
        return {"target_type": old["target_type"], "target_id": old["target_id"],
                "restored_from_rev": old["rev"], "new_rev": new_rev}


# ---------------------------------------------------------------- comments

@mcp.tool()
def comment_create(target_type: str, target_id: str, body: str,
                   anchor_quote: str = "", parent_id: str | None = None) -> dict:
    """Comment on an entity or chapter. anchor_quote pins the comment to a quoted span of
    the current content; parent_id makes it a threaded reply. Any role."""
    with _db() as conn:
        target = _get_target(conn, target_type, target_id)
        if anchor_quote and anchor_quote not in (target["content_md"] or ""):
            raise ValueError("anchor_quote not found in current content — quote it exactly")
        if parent_id is not None:
            parent = conn.execute(
                "SELECT target_type, target_id FROM comments WHERE id=?",
                (parent_id,)).fetchone()
            if parent is None:
                raise ValueError(f"parent comment {parent_id} not found")
            if parent["target_type"] != target_type or parent["target_id"] != target_id:
                raise ValueError("parent comment is on a different record — replies must "
                                 "stay on the same entity/chapter/scene")
        cid = _id()
        conn.execute(
            "INSERT INTO comments (id, target_type, target_id, parent_id, anchor_quote, body, "
            "created_by, created_at) VALUES (?,?,?,?,?,?,?,?)",
            (cid, target_type, target_id, parent_id, anchor_quote, body, _caller()["name"], _now()),
        )
        return _row(conn.execute("SELECT * FROM comments WHERE id=?", (cid,)).fetchone())


@mcp.tool()
def comment_list(target_type: str, target_id: str, status: str | None = None) -> list[dict]:
    """List comments on an entity or chapter, oldest first. status: open | resolved | None for all."""
    with _db() as conn:
        _get_target(conn, target_type, target_id)
        q = "SELECT * FROM comments WHERE target_type=? AND target_id=?"
        args: list = [target_type, target_id]
        if status is not None:
            q += " AND status=?"
            args.append(status)
        return _rows(conn.execute(q + " ORDER BY created_at", args))


@mcp.tool()
def comment_resolve(comment_id: str) -> dict:
    """Mark a comment resolved. Any role."""
    with _db() as conn:
        if conn.execute("SELECT 1 FROM comments WHERE id=?", (comment_id,)).fetchone() is None:
            raise ValueError(f"comment {comment_id} not found")
        conn.execute("UPDATE comments SET status='resolved', resolved_by=?, resolved_at=? WHERE id=?",
                     (_caller()["name"], _now(), comment_id))
        return _row(conn.execute("SELECT * FROM comments WHERE id=?", (comment_id,)).fetchone())


# ---------------------------------------------------------------- proposals

@mcp.tool()
def proposal_create(target_type: str, target_id: str, proposed_content_md: str,
                    rationale: str = "") -> dict:
    """Propose replacement content for an entity or chapter (track-changes style).
    Recorded against the target's current revision; an author accepts or rejects it.
    Any role — this is THE write path for editor keys."""
    with _db() as conn:
        target = _get_target(conn, target_type, target_id)
        pid = _id()
        conn.execute(
            "INSERT INTO proposals (id, target_type, target_id, base_rev, proposed_content_md, "
            "rationale, created_by, created_at) VALUES (?,?,?,?,?,?,?,?)",
            (pid, target_type, target_id, target["rev"], proposed_content_md, rationale,
             _caller()["name"], _now()),
        )
        return _row(conn.execute("SELECT * FROM proposals WHERE id=?", (pid,)).fetchone())


@mcp.tool()
def proposal_list(status: str = "pending", target_type: str | None = None,
                  target_id: str | None = None) -> list[dict]:
    """List proposals (default: pending), optionally scoped to one target. Content omitted;
    each row includes 'stale' = base_rev is behind the target's current rev."""
    with _db() as conn:
        q = ("SELECT id, target_type, target_id, base_rev, rationale, status, created_by, "
             "created_at, decided_by, decided_at, length(proposed_content_md) AS content_chars "
             "FROM proposals WHERE 1=1")
        args: list = []
        if status != "all":
            q += " AND status=?"
            args.append(status)
        if target_type is not None and target_id is not None:
            q += " AND target_type=? AND target_id=?"
            args += [target_type, target_id]
        out = _rows(conn.execute(q + " ORDER BY created_at DESC", args))
        for p in out:
            table = TARGET_TABLES[p["target_type"]]
            cur = conn.execute(f"SELECT rev FROM {table} WHERE id=?", (p["target_id"],)).fetchone()
            p["stale"] = bool(cur) and cur["rev"] > p["base_rev"]
        return out


@mcp.tool()
def proposal_get(proposal_id: str) -> dict:
    """Get a proposal with full proposed content plus the target's CURRENT content for diffing."""
    with _db() as conn:
        p = _row(conn.execute("SELECT * FROM proposals WHERE id=?", (proposal_id,)).fetchone())
        if p is None:
            raise ValueError(f"proposal {proposal_id} not found")
        target = _get_target(conn, p["target_type"], p["target_id"])
        p["current_rev"] = target["rev"]
        p["current_content_md"] = target["content_md"]
        p["stale"] = target["rev"] > p["base_rev"]
        return p


@mcp.tool()
def proposal_accept(proposal_id: str, force: bool = False, note: str = "") -> dict:
    """Accept a pending proposal: its content becomes a new revision on the target,
    attributed to the proposer. Refuses stale proposals (target moved since base_rev)
    unless force=true. Author role required."""
    _require_author()
    who = _caller()["name"]
    with _db() as conn:
        p = conn.execute("SELECT * FROM proposals WHERE id=?", (proposal_id,)).fetchone()
        if p is None:
            raise ValueError(f"proposal {proposal_id} not found")
        if p["status"] != "pending":
            raise ValueError(f"proposal is already {p['status']}")
        target = _get_target(conn, p["target_type"], p["target_id"])
        _require_unlocked(conn, p["target_type"], p["target_id"], "proposal_accept")
        if target["rev"] > p["base_rev"] and not force:
            raise ValueError(
                f"stale: proposal was written against rev {p['base_rev']} but target is at "
                f"rev {target['rev']}. Review with proposal_get, then pass force=true to apply anyway."
            )
        new_rev = _write_revision(
            conn, p["target_type"], p["target_id"], p["proposed_content_md"],
            f"proposal {proposal_id} by {p['created_by']}, accepted by {who}"
            + (f": {note}" if note else ""),
            p["created_by"],
        )
        conn.execute(
            "UPDATE proposals SET status='accepted', decided_by=?, decided_at=?, decision_note=? WHERE id=?",
            (who, _now(), note, proposal_id),
        )
        return {"proposal_id": proposal_id, "applied_as_rev": new_rev,
                "target_type": p["target_type"], "target_id": p["target_id"]}


@mcp.tool()
def proposal_reject(proposal_id: str, note: str = "") -> dict:
    """Reject a pending proposal with an optional note for the proposer. Author role required."""
    _require_author()
    with _db() as conn:
        p = conn.execute("SELECT * FROM proposals WHERE id=?", (proposal_id,)).fetchone()
        if p is None:
            raise ValueError(f"proposal {proposal_id} not found")
        if p["status"] != "pending":
            raise ValueError(f"proposal is already {p['status']}")
        conn.execute(
            "UPDATE proposals SET status='rejected', decided_by=?, decided_at=?, decision_note=? WHERE id=?",
            (_caller()["name"], _now(), note, proposal_id),
        )
        return _row(conn.execute("SELECT * FROM proposals WHERE id=?", (proposal_id,)).fetchone())


# ---------------------------------------------------------------- search

def _search_like(conn, query: str, project_id: str | None, limit: int) -> list[dict]:
    """Substring fallback for queries FTS can't parse or doesn't match."""
    like = f"%{query}%"
    out: list[dict] = []
    eq = ("SELECT id, project_id, kind, name, summary, content_md FROM entities "
          "WHERE deleted=0 AND (name LIKE ? OR summary LIKE ? OR content_md LIKE ?)")
    cq = ("SELECT id, project_id, title, content_md FROM chapters "
          "WHERE deleted=0 AND (title LIKE ? OR content_md LIKE ?)")
    sq = ("SELECT id, project_id, title, synopsis, content_md FROM scenes "
          "WHERE deleted=0 AND (title LIKE ? OR synopsis LIKE ? OR content_md LIKE ?)")
    eargs: list = [like, like, like]
    cargs: list = [like, like]
    sargs: list = [like, like, like]
    if project_id is not None:
        eq += " AND project_id=?"
        cq += " AND project_id=?"
        sq += " AND project_id=?"
        eargs.append(project_id)
        cargs.append(project_id)
        sargs.append(project_id)
    ql = query.lower()
    for r in conn.execute(eq + " LIMIT ?", eargs + [limit]):
        text = r["content_md"] or r["summary"] or ""
        i = text.lower().find(ql)
        out.append({"type": "entity", "id": r["id"], "project_id": r["project_id"],
                    "kind": r["kind"], "name": r["name"],
                    "snippet": text[max(0, i - 60):i + 120] if i >= 0 else text[:120]})
    for r in conn.execute(cq + " LIMIT ?", cargs + [limit]):
        text = r["content_md"] or ""
        i = text.lower().find(ql)
        out.append({"type": "chapter", "id": r["id"], "project_id": r["project_id"],
                    "name": r["title"],
                    "snippet": text[max(0, i - 60):i + 120] if i >= 0 else text[:120]})
    for r in conn.execute(sq + " LIMIT ?", sargs + [limit]):
        text = r["content_md"] or r["synopsis"] or ""
        i = text.lower().find(ql)
        out.append({"type": "scene", "id": r["id"], "project_id": r["project_id"],
                    "name": r["title"],
                    "snippet": text[max(0, i - 60):i + 120] if i >= 0 else text[:120]})
    return out[:limit]


@mcp.tool()
def search(query: str, project_id: str | None = None, types: str = "",
           limit: int = 20) -> list[dict]:
    """Ranked full-text search (FTS5, stemmed) across entities, chapters, and scenes.
    types: optional comma list to restrict, e.g. "scene,chapter". Falls back to
    substring matching when FTS has no hits. Returns matches with snippets."""
    type_filter = {t.strip() for t in types.split(",") if t.strip()}
    if type_filter - set(TARGET_TABLES):
        raise ValueError(f"types must be from {sorted(TARGET_TABLES)}")
    # Quote each term so user text can't break FTS query syntax; AND semantics.
    terms = [t.replace('"', '""') for t in query.split() if t]
    if not terms:
        return []
    match = " ".join(f'"{t}"' for t in terms)
    out: list[dict] = []
    with _db() as conn:
        q = ("SELECT target_type, target_id, project_id, name, "
             "snippet(fts, 4, '', '', '…', 24) AS snip, bm25(fts) AS rank "
             "FROM fts WHERE fts MATCH ?")
        args: list = [match]
        if project_id is not None:
            q += " AND project_id=?"
            args.append(project_id)
        try:
            rows = conn.execute(q + " ORDER BY rank LIMIT ?", args + [limit * 2]).fetchall()
        except sqlite3.OperationalError:
            rows = []
        for r in rows:
            if type_filter and r["target_type"] not in type_filter:
                continue
            hit = {"type": r["target_type"], "id": r["target_id"],
                   "project_id": r["project_id"], "name": r["name"].strip(),
                   "snippet": r["snip"]}
            if r["target_type"] == "entity":
                ent = conn.execute("SELECT kind, name FROM entities WHERE id=?",
                                   (r["target_id"],)).fetchone()
                if ent is not None:
                    hit["kind"] = ent["kind"]
                    hit["name"] = ent["name"]
            out.append(hit)
            if len(out) >= limit:
                break
        if not out and not type_filter:
            out = _search_like(conn, query, project_id, limit)
    return out


# ---------------------------------------------------------------- deterministic gates

def _profile_for(conn, target_type: str, target_id: str) -> dict | None:
    """Resolve the voice profile for a node: its own voice_profile_id meta, else its
    chapter's (for scenes). Returns {'id', 'banned_chars', 'patterns'} or None."""
    pid = conn.execute(
        "SELECT value FROM node_meta WHERE target_type=? AND target_id=? AND key='voice_profile_id'",
        (target_type, target_id)).fetchone()
    if pid is None and target_type == "scene":
        sc = conn.execute("SELECT chapter_id FROM scenes WHERE id=?", (target_id,)).fetchone()
        if sc is not None:
            pid = conn.execute(
                "SELECT value FROM node_meta WHERE target_type='chapter' AND target_id=? "
                "AND key='voice_profile_id'", (sc["chapter_id"],)).fetchone()
    if pid is None:
        return None
    meta = {r["key"]: r["value"] for r in conn.execute(
        "SELECT key, value FROM node_meta WHERE target_type='entity' AND target_id=?",
        (pid["value"],))}
    try:
        patterns = json.loads(meta.get("lint_banned_patterns", "[]"))
    except ValueError:
        patterns = []
    return {"id": pid["value"], "banned_chars": meta.get("lint_banned", ""),
            "patterns": patterns}


def _lint_violations(conn, target_type: str, target_id: str) -> list[dict]:
    """Profile-scoped mechanical style check (spec B): banned characters/patterns
    attached to the node's voice profile. No profile or no rules = no violations."""
    prof = _profile_for(conn, target_type, target_id)
    if not prof or (not prof["banned_chars"] and not prof["patterns"]):
        return []
    row = conn.execute(f"SELECT content_md FROM {TARGET_TABLES[target_type]} WHERE id=?",
                       (target_id,)).fetchone()
    text = (row["content_md"] or "") if row else ""
    out = []
    for i, para in enumerate(text.split("\n\n"), start=1):
        for ch in prof["banned_chars"]:
            n = para.count(ch)
            if n:
                out.append({"check": "voice_lint", "paragraph": i, "char": ch, "count": n,
                            "snippet": para.strip()[:80]})
        for pat in prof["patterns"]:
            try:
                if re.search(pat, para):
                    out.append({"check": "voice_lint", "paragraph": i, "pattern": pat,
                                "snippet": para.strip()[:80]})
            except re.error:
                continue
    return out


def _silhouette_terms(conn, project_id: str) -> dict[str, list[str]]:
    """entity_id -> word-bounded terms that must never appear in prose."""
    out: dict[str, list[str]] = {}
    for r in conn.execute(
            "SELECT id, name FROM entities WHERE project_id=? AND visibility='silhouette' "
            "AND deleted=0", (project_id,)):
        terms = [r["name"]]
        for key in ("aliases", "defined_phrases"):
            m = conn.execute(
                "SELECT value FROM node_meta WHERE target_type='entity' AND target_id=? AND key=?",
                (r["id"], key)).fetchone()
            if m is not None:
                terms += [t.strip() for t in m["value"].split(",")]
        out[r["id"]] = [t for t in terms if len(t.strip()) >= 3]
    return out


def _silhouette_leaks_for(conn, target_type: str, target_id: str) -> list[dict]:
    row = conn.execute(
        f"SELECT project_id, content_md FROM {TARGET_TABLES[target_type]} WHERE id=?",
        (target_id,)).fetchone()
    if row is None or not row["content_md"]:
        return []
    text = row["content_md"].lower()
    leaks = []
    for eid, terms in _silhouette_terms(conn, row["project_id"]).items():
        for t in terms:
            for m in re.finditer(rf"(?<!\w){re.escape(t.lower())}(?!\w)", text):
                leaks.append({"check": "silhouette_leak", "entity_id": eid, "term": t,
                              "snippet": row["content_md"][max(0, m.start() - 40):m.end() + 40]})
    return leaks


def _gate_final(conn, target_type: str, target_id: str, final_override: bool):
    """Deterministic final gate (spec H.5 / F / B): voice lint + silhouette leaks +
    linked claims verified false. Override is owner-only and logs a decision."""
    problems = _lint_violations(conn, target_type, target_id)
    problems += _silhouette_leaks_for(conn, target_type, target_id)
    for l in conn.execute("SELECT from_id, to_id FROM links WHERE from_id=? OR to_id=?",
                          (target_id, target_id)):
        other = l["to_id"] if l["from_id"] == target_id else l["from_id"]
        v = conn.execute("SELECT verdict FROM verifications WHERE claim_id=? "
                         "ORDER BY created_at DESC LIMIT 1", (other,)).fetchone()
        if v is not None and v["verdict"] == "false":
            ent = conn.execute("SELECT name FROM entities WHERE id=?", (other,)).fetchone()
            problems.append({"check": "fact", "claim_id": other,
                             "detail": f"linked claim verified FALSE: {ent['name'] if ent else other}"})
    if not problems:
        return
    if not final_override:
        raise ValueError(
            f"final gate BLOCKED ({len(problems)} problem(s)): "
            + json.dumps(problems)[:1400]
            + " — fix these, or an owner passes final_override=true (logs a decision).")
    _require_owner()
    row = conn.execute(f"SELECT project_id FROM {TARGET_TABLES[target_type]} WHERE id=?",
                       (target_id,)).fetchone()
    conn.execute(
        "INSERT INTO decisions (id, project_id, subject_type, subject_id, ruling, rationale, "
        "decided_by, created_at) VALUES (?,?,?,?,?,?,?,?)",
        (_id(), row["project_id"], "gate_override", target_id,
         f"{target_type} allowed to enter final despite gate report",
         json.dumps(problems)[:2000], _caller()["name"], _now()))


@mcp.tool()
def voice_lint_run(target_type: str, target_id: str) -> list[dict]:
    """Run the profile-scoped mechanical style check on a chapter or scene. Attach a
    profile via meta key 'voice_profile_id' pointing at a voice_profile entity whose
    meta 'lint_banned' lists banned characters (e.g. em dash, semicolon, colon) and
    optional 'lint_banned_patterns' (JSON regex list). Any role."""
    if target_type not in ("chapter", "scene"):
        raise ValueError("target_type must be 'chapter' or 'scene'")
    with _db() as conn:
        _get_target(conn, target_type, target_id)
        return _lint_violations(conn, target_type, target_id)


@mcp.tool()
def silhouette_leak_check(project_id: str, chapter_id: str | None = None) -> list[dict]:
    """Scan prose for silhouette-visibility entities' names/aliases/defined phrases
    (meta 'defined_phrases', comma-separated). A leak is blocking at the final gate.
    Scans one chapter (+its scenes) or the whole project. Any role."""
    leaks: list[dict] = []
    with _db() as conn:
        if chapter_id is not None:
            targets = [("chapter", chapter_id)] + [
                ("scene", r["id"]) for r in conn.execute(
                    "SELECT id FROM scenes WHERE chapter_id=? AND deleted=0", (chapter_id,))]
        else:
            targets = [("chapter", r["id"]) for r in conn.execute(
                "SELECT id FROM chapters WHERE project_id=? AND deleted=0", (project_id,))]
            targets += [("scene", r["id"]) for r in conn.execute(
                "SELECT id FROM scenes WHERE project_id=? AND deleted=0", (project_id,))]
        for tt, tid in targets:
            for leak in _silhouette_leaks_for(conn, tt, tid):
                leaks.append({**leak, "target_type": tt, "target_id": tid})
    return leaks


# ---------------------------------------------------------------- parts

@mcp.tool()
def part_create(project_id: str, title: str, sort_order: int = 0) -> dict:
    """Create a manuscript part (Part → Chapter → Scene). Author role required."""
    _require_author()
    with _db() as conn:
        if conn.execute("SELECT 1 FROM projects WHERE id=? AND deleted=0",
                        (project_id,)).fetchone() is None:
            raise ValueError(f"project {project_id} not found")
        pid = _id()
        conn.execute("INSERT INTO parts (id, project_id, title, sort_order, created_by, created_at) "
                     "VALUES (?,?,?,?,?,?)",
                     (pid, project_id, title, sort_order, _caller()["name"], _now()))
        return _row(conn.execute("SELECT * FROM parts WHERE id=?", (pid,)).fetchone())


@mcp.tool()
def part_list(project_id: str) -> list[dict]:
    """List a project's parts in order, with chapter counts."""
    with _db() as conn:
        out = _rows(conn.execute(
            "SELECT * FROM parts WHERE project_id=? AND deleted=0 ORDER BY sort_order, created_at",
            (project_id,)))
        for p in out:
            p["chapter_count"] = conn.execute(
                "SELECT COUNT(*) n FROM chapters WHERE part_id=? AND deleted=0",
                (p["id"],)).fetchone()["n"]
        return out


@mcp.tool()
def part_update(part_id: str, title: str | None = None, sort_order: int | None = None) -> dict:
    """Rename/reorder a part. Author role required."""
    _require_author()
    with _db() as conn:
        if conn.execute("SELECT 1 FROM parts WHERE id=? AND deleted=0", (part_id,)).fetchone() is None:
            raise ValueError(f"part {part_id} not found")
        for field, val in (("title", title), ("sort_order", sort_order)):
            if val is not None:
                conn.execute(f"UPDATE parts SET {field}=? WHERE id=?", (val, part_id))
        return _row(conn.execute("SELECT * FROM parts WHERE id=?", (part_id,)).fetchone())


# ---------------------------------------------------------------- locks & visibility

@mcp.tool()
def lock_set(target_type: str, target_id: str, kind: str = "content", reason: str = "") -> dict:
    """Lock an entity/chapter/scene against ALL mutation (kind: content |
    personal_truth). Blocked attempts are recorded. Owner only."""
    _require_owner()
    if kind not in ("content", "personal_truth"):
        raise ValueError("kind must be 'content' or 'personal_truth'")
    with _db() as conn:
        _get_target(conn, target_type, target_id)
        conn.execute("INSERT OR REPLACE INTO locks (target_type, target_id, kind, reason, "
                     "locked_by, locked_at) VALUES (?,?,?,?,?,?)",
                     (target_type, target_id, kind, reason, _caller()["name"], _now()))
        _log_lock_event(conn, target_type, target_id, "locked", reason)
        return {"locked": target_id, "kind": kind}


@mcp.tool()
def lock_remove(target_type: str, target_id: str) -> dict:
    """Remove a lock. Owner only."""
    _require_owner()
    with _db() as conn:
        n = conn.execute("DELETE FROM locks WHERE target_type=? AND target_id=?",
                         (target_type, target_id)).rowcount
        _log_lock_event(conn, target_type, target_id, "unlocked")
        return {"unlocked": target_id, "existed": bool(n)}


@mcp.tool()
def lock_list(include_events: bool = False) -> dict:
    """List all locks (and optionally the lock event log, newest first)."""
    with _db() as conn:
        out = {"locks": _rows(conn.execute("SELECT * FROM locks ORDER BY locked_at"))}
        if include_events:
            out["events"] = _rows(conn.execute(
                "SELECT * FROM lock_events ORDER BY created_at DESC LIMIT 200"))
        return out


@mcp.tool()
def visibility_set(entity_id: str, visibility: str) -> dict:
    """Set canon visibility: public | private | silhouette. Silhouette facts drive
    continuity but must never be named in prose (enforced at the final gate).
    Relaxing private/silhouette back to public is owner-only. Author role required."""
    _require_author()
    if visibility not in VISIBILITIES:
        raise ValueError(f"visibility must be one of {sorted(VISIBILITIES)}")
    with _db() as conn:
        ent = _get_target(conn, "entity", entity_id)
        if ent["visibility"] in ("private", "silhouette") and visibility == "public":
            _require_owner()
        conn.execute("UPDATE entities SET visibility=?, updated_at=? WHERE id=?",
                     (visibility, _now(), entity_id))
        return {"entity_id": entity_id, "visibility": visibility}


# ---------------------------------------------------------------- seam pointer

@mcp.tool()
def seam_set(project_id: str, last_verified: str, first_invented: str,
             seam_date: str = "") -> dict:
    """Set the reality→fiction seam: the last verified real-world event and the first
    invented one (spec B). seam_date (ISO) enables mechanical pre/post sorting.
    Author role required."""
    _require_author()
    with _db() as conn:
        if conn.execute("SELECT 1 FROM projects WHERE id=? AND deleted=0",
                        (project_id,)).fetchone() is None:
            raise ValueError(f"project {project_id} not found")
        who = _caller()["name"]
        for k, v in (("seam_last_verified", last_verified),
                     ("seam_first_invented", first_invented), ("seam_date", seam_date)):
            conn.execute(
                "INSERT INTO node_meta (target_type, target_id, key, value, updated_by, updated_at) "
                "VALUES ('project',?,?,?,?,?) ON CONFLICT(target_type, target_id, key) "
                "DO UPDATE SET value=excluded.value, updated_by=excluded.updated_by, "
                "updated_at=excluded.updated_at", (project_id, k, v, who, _now()))
        # read back on THIS conn — a fresh connection wouldn't see the open transaction
        return {"last_verified": last_verified, "first_invented": first_invented,
                "seam_date": seam_date}


@mcp.tool()
def seam_get(project_id: str) -> dict:
    """Read the reality→fiction seam pointer."""
    with _db() as conn:
        m = {r["key"]: r["value"] for r in conn.execute(
            "SELECT key, value FROM node_meta WHERE target_type='project' AND target_id=? "
            "AND key IN ('seam_last_verified','seam_first_invented','seam_date')",
            (project_id,))}
        return {"last_verified": m.get("seam_last_verified", ""),
                "first_invented": m.get("seam_first_invented", ""),
                "seam_date": m.get("seam_date", "")}


# ---------------------------------------------------------------- research claims & verification

@mcp.tool()
def research_claim_create(project_id: str, claim: str, domain: str,
                          classification: str, applicable_date: str = "",
                          asserted_by_character: str = "", source_url: str = "",
                          notes: str = "") -> dict:
    """Log a factual assertion as a research claim (spec H). domain: technical |
    scientific | political | historical | geographic | medical | legal | economic |
    cultural | other. classification: fact | inference | projection |
    deliberate_fiction. Link it to the scenes/chapters that assert it via
    link_create(rel_type='asserts'). Author role required."""
    _require_author()
    if domain not in CLAIM_DOMAINS:
        raise ValueError(f"domain must be one of {sorted(CLAIM_DOMAINS)}")
    if classification not in CLAIM_CLASSES:
        raise ValueError(f"classification must be one of {sorted(CLAIM_CLASSES)}")
    who = _caller()["name"]
    with _db() as conn:
        if conn.execute("SELECT 1 FROM projects WHERE id=? AND deleted=0",
                        (project_id,)).fetchone() is None:
            raise ValueError(f"project {project_id} not found")
        eid = _id()
        conn.execute(
            "INSERT INTO entities (id, project_id, kind, name, summary, content_md, sort_order, "
            "created_by, created_at, updated_at) VALUES (?,?,?,?,?,?,0,?,?,?)",
            (eid, project_id, "research", claim[:120], f"[{domain}/{classification}]",
             claim + (f"\n\n{notes}" if notes else ""), who, _now(), _now()))
        _record_initial_revision(conn, "entity", eid, claim, who)
        seam_date = conn.execute(
            "SELECT value FROM node_meta WHERE target_type='project' AND target_id=? "
            "AND key='seam_date'", (project_id,)).fetchone()
        seam_side = ""
        if seam_date is not None and seam_date["value"] and applicable_date:
            seam_side = "pre" if applicable_date <= seam_date["value"] else "post"
        for k, v in (("claim_domain", domain), ("claim_class", classification),
                     ("applicable_date", applicable_date),
                     ("asserted_by_character", asserted_by_character),
                     ("source_url", source_url), ("seam_side", seam_side)):
            if v:
                conn.execute(
                    "INSERT INTO node_meta (target_type, target_id, key, value, updated_by, updated_at) "
                    "VALUES ('entity',?,?,?,?,?)", (eid, k, v, who, _now()))
        return _row(conn.execute("SELECT * FROM entities WHERE id=?", (eid,)).fetchone())


@mcp.tool()
def research_claim_verify(claim_id: str, verdict: str, sources: list[dict] | None = None,
                          confidence: float = 0.0, notes: str = "") -> dict:
    """File an immutable, source-cited verification of a claim. verdict: verified |
    false | disputed | unverifiable | outdated. sources: [{url, title, publisher,
    access_date, quote, source_type(primary|secondary|commentary)}] — required for
    verified/false verdicts. Any role (evidence-gathering, not editorial judgment)."""
    if verdict not in VERDICTS:
        raise ValueError(f"verdict must be one of {sorted(VERDICTS)}")
    sources = sources or []
    if verdict in ("verified", "false") and not sources:
        raise ValueError(f"a '{verdict}' verdict requires at least one cited source")
    who = _caller()["name"]
    with _db() as conn:
        claim = _get_target(conn, "entity", claim_id)
        vid = _id()
        conn.execute(
            "INSERT INTO verifications (id, claim_id, claim_rev, verdict, confidence, "
            "sources_json, notes, created_by, created_at) VALUES (?,?,?,?,?,?,?,?,?)",
            (vid, claim_id, claim["rev"], verdict, confidence, json.dumps(sources),
             notes, who, _now()))
        conn.execute(
            "INSERT INTO node_meta (target_type, target_id, key, value, updated_by, updated_at) "
            "VALUES ('entity',?,?,?,?,?) ON CONFLICT(target_type, target_id, key) "
            "DO UPDATE SET value=excluded.value, updated_by=excluded.updated_by, "
            "updated_at=excluded.updated_at",
            (claim_id, "verification_status", verdict, who, _now()))
        return _row(conn.execute("SELECT * FROM verifications WHERE id=?", (vid,)).fetchone())


@mcp.tool()
def research_claim_verifications(claim_id: str) -> list[dict]:
    """Verification history for a claim, newest first, sources parsed."""
    with _db() as conn:
        _get_target(conn, "entity", claim_id)
        out = _rows(conn.execute(
            "SELECT * FROM verifications WHERE claim_id=? ORDER BY created_at DESC",
            (claim_id,)))
        for v in out:
            try:
                v["sources"] = json.loads(v.pop("sources_json"))
            except ValueError:
                v["sources"] = []
        return out


# ---------------------------------------------------------------- fictionalization, decisions, rebuttals

@mcp.tool()
def fictionalization_log_create(project_id: str, real_fact: str, invented_fact: str,
                                rationale: str = "", target_type: str = "",
                                target_id: str = "") -> dict:
    """Record a deliberate fictionalization: the real fact and the invention that
    replaces it in prose, so inventions never harden into remembered biography
    (spec 3.12). Author role required."""
    _require_author()
    with _db() as conn:
        fid = _id()
        conn.execute(
            "INSERT INTO fictionalizations (id, project_id, real_fact, invented_fact, rationale, "
            "target_type, target_id, created_by, created_at) VALUES (?,?,?,?,?,?,?,?,?)",
            (fid, project_id, real_fact, invented_fact, rationale, target_type, target_id,
             _caller()["name"], _now()))
        return _row(conn.execute("SELECT * FROM fictionalizations WHERE id=?", (fid,)).fetchone())


@mcp.tool()
def fictionalization_log_list(project_id: str) -> list[dict]:
    """List a project's fictionalization records."""
    with _db() as conn:
        return _rows(conn.execute(
            "SELECT * FROM fictionalizations WHERE project_id=? ORDER BY created_at",
            (project_id,)))


@mcp.tool()
def decision_create(project_id: str, subject_type: str, ruling: str,
                    subject_id: str = "", rationale: str = "") -> dict:
    """Record an explicit owner decision (spec 3.13): the ruling of record on a
    finding, proposal, claim, gate, or canon question. Immutable. Owner only."""
    _require_owner()
    with _db() as conn:
        did = _id()
        conn.execute(
            "INSERT INTO decisions (id, project_id, subject_type, subject_id, ruling, rationale, "
            "decided_by, created_at) VALUES (?,?,?,?,?,?,?,?)",
            (did, project_id, subject_type, subject_id, ruling, rationale,
             _caller()["name"], _now()))
        return _row(conn.execute("SELECT * FROM decisions WHERE id=?", (did,)).fetchone())


@mcp.tool()
def decision_list(project_id: str, subject_id: str | None = None) -> list[dict]:
    """List decisions for a project (optionally scoped to one subject), newest first."""
    with _db() as conn:
        q = "SELECT * FROM decisions WHERE project_id=?"
        args: list = [project_id]
        if subject_id is not None:
            q += " AND subject_id=?"
            args.append(subject_id)
        return _rows(conn.execute(q + " ORDER BY created_at DESC", args))


@mcp.tool()
def rebuttal_create(target_kind: str, target_id: str, body: str,
                    evidence_quote: str = "", location: str = "") -> dict:
    """File structured dissent against a proposal (or, later, a finding) — same
    evidence ethic as findings (spec 3.13). The dissent trail survives the decision
    either way. Any role."""
    if target_kind not in ("proposal", "finding"):
        raise ValueError("target_kind must be 'proposal' or 'finding'")
    with _db() as conn:
        if target_kind == "proposal" and conn.execute(
                "SELECT 1 FROM proposals WHERE id=?", (target_id,)).fetchone() is None:
            raise ValueError(f"proposal {target_id} not found")
        if target_kind == "finding" and conn.execute(
                "SELECT 1 FROM findings WHERE id=?", (target_id,)).fetchone() is None:
            raise ValueError(f"finding {target_id} not found")
        rid = _id()
        conn.execute(
            "INSERT INTO rebuttals (id, target_kind, target_id, body, evidence_quote, location, "
            "created_by, created_at) VALUES (?,?,?,?,?,?,?,?)",
            (rid, target_kind, target_id, body, evidence_quote, location,
             _caller()["name"], _now()))
        return _row(conn.execute("SELECT * FROM rebuttals WHERE id=?", (rid,)).fetchone())


@mcp.tool()
def rebuttal_list(target_kind: str, target_id: str) -> list[dict]:
    """List rebuttals filed against a proposal or finding, oldest first."""
    with _db() as conn:
        return _rows(conn.execute(
            "SELECT * FROM rebuttals WHERE target_kind=? AND target_id=? ORDER BY created_at",
            (target_kind, target_id)))


# ---------------------------------------------------------------- editorial engine (Sol)

SOL_OUTPUT_SCHEMA = {
    "analysis_type": "chapter_gate | scene_check | ...",
    "target_id": "id", "target_revision": 0,
    "model": "gpt-5.6-sol", "reasoning_effort": "max",
    "verdict": "pass | revise",
    "intent": {"source": "accepted | inferred", "summary": "..."},
    "observed": {"summary": "..."},
    "strengths_to_protect": [{"evidence_quote": "EXACT quotation", "location": "…",
                              "explanation": "why protect"}],
    "scores": {"dimension_name": 0},
    "findings": [{"severity": "blocking | major | minor | watch", "category": "…",
                  "confidence": 0.0, "evidence_quote": "EXACT quotation", "location": "…",
                  "explanation": "why this is a story problem",
                  "affected_entity_ids": [], "smallest_intervention": "…"}],
    "limitations": [],
}


def _rev_content(conn, target_type: str, target_id: str, rev: int) -> str:
    r = conn.execute(
        "SELECT content_md FROM revisions WHERE target_type=? AND target_id=? AND rev=?",
        (target_type, target_id, rev)).fetchone()
    return r["content_md"] if r is not None else ""


def _run_project(conn, target_type: str, target_id: str) -> str:
    row = conn.execute(f"SELECT project_id FROM {TARGET_TABLES[target_type]} WHERE id=?",
                       (target_id,)).fetchone()
    if row is None:
        raise ValueError(f"{target_type} {target_id} not found")
    return row["project_id"]


@mcp.tool()
def analysis_run_create(analysis_type: str, target_type: str, target_id: str,
                        model: str = "gpt-5.6-sol", reasoning_effort: str = "max") -> dict:
    """Queue a formal editorial analysis (spec §10). Pins the target revision and the
    applicable profile revisions for calibration. A runner picks up queued jobs,
    assembles the packet via review_packet_get, runs the model, and posts results via
    analysis_run_complete. second_opinion and cold_read are owner-triggered pull-cords
    and always advisory. Any role may request standard reviews."""
    if analysis_type not in ANALYSIS_TYPES:
        raise ValueError(f"analysis_type must be one of {sorted(ANALYSIS_TYPES)}")
    if analysis_type in ADVISORY_TYPES:
        _require_owner()
    with _db() as conn:
        target = _get_target(conn, target_type, target_id)
        project_id = target["project_id"]
        pinned = {}
        for mt_type, mt_id, key in (
                (target_type, target_id, "voice_profile_id"),
                (target_type, target_id, "rubric_profile_id")):
            m = conn.execute(
                "SELECT value FROM node_meta WHERE target_type=? AND target_id=? AND key=?",
                (mt_type, mt_id, key)).fetchone()
            if m is not None:
                prof = conn.execute("SELECT rev FROM entities WHERE id=?",
                                    (m["value"],)).fetchone()
                if prof is not None:
                    pinned[key] = {"id": m["value"], "rev": prof["rev"]}
        rid = _id()
        conn.execute(
            "INSERT INTO analysis_runs (id, project_id, analysis_type, target_type, target_id, "
            "target_rev, pinned_json, status, model, reasoning_effort, advisory, requested_by, "
            "created_at) VALUES (?,?,?,?,?,?,?,'queued',?,?,?,?,?)",
            (rid, project_id, analysis_type, target_type, target_id, target["rev"],
             json.dumps(pinned), model, reasoning_effort,
             int(analysis_type in ADVISORY_TYPES), _caller()["name"], _now()))
        return _row(conn.execute("SELECT * FROM analysis_runs WHERE id=?", (rid,)).fetchone())


@mcp.tool()
def analysis_run_list(project_id: str | None = None, status: str | None = None) -> list[dict]:
    """List analysis runs (newest first), optionally by project and/or status
    (queued | running | complete | cancelled). Each includes a computed 'stale' flag —
    true when the target has moved past the pinned revision."""
    with _db() as conn:
        q = "SELECT * FROM analysis_runs WHERE 1=1"
        args: list = []
        if project_id is not None:
            q += " AND project_id=?"
            args.append(project_id)
        if status is not None:
            q += " AND status=?"
            args.append(status)
        out = _rows(conn.execute(q + " ORDER BY created_at DESC LIMIT 100", args))
        for r in out:
            cur = conn.execute(
                f"SELECT rev FROM {TARGET_TABLES[r['target_type']]} WHERE id=?",
                (r["target_id"],)).fetchone()
            r["stale"] = bool(cur) and cur["rev"] > r["target_rev"]
        return out


@mcp.tool()
def analysis_run_get(run_id: str) -> dict:
    """Get an analysis run with its findings and strengths, scores parsed, staleness computed."""
    with _db() as conn:
        r = _row(conn.execute("SELECT * FROM analysis_runs WHERE id=?", (run_id,)).fetchone())
        if r is None:
            raise ValueError(f"analysis run {run_id} not found")
        r["scores"] = json.loads(r.pop("scores_json") or "{}")
        r["limitations"] = json.loads(r.pop("limitations_json") or "[]")
        r["pinned"] = json.loads(r.pop("pinned_json") or "{}")
        cur = conn.execute(f"SELECT rev FROM {TARGET_TABLES[r['target_type']]} WHERE id=?",
                           (r["target_id"],)).fetchone()
        r["stale"] = bool(cur) and cur["rev"] > r["target_rev"]
        r["findings"] = _rows(conn.execute(
            "SELECT * FROM findings WHERE run_id=? ORDER BY created_at", (run_id,)))
        r["strengths"] = _rows(conn.execute(
            "SELECT * FROM strengths WHERE run_id=? ORDER BY created_at", (run_id,)))
        return r


@mcp.tool()
def analysis_run_cancel(run_id: str) -> dict:
    """Cancel a queued or running analysis. Any role."""
    with _db() as conn:
        r = conn.execute("SELECT status FROM analysis_runs WHERE id=?", (run_id,)).fetchone()
        if r is None:
            raise ValueError(f"analysis run {run_id} not found")
        if r["status"] not in ("queued", "running"):
            raise ValueError(f"run is already {r['status']}")
        conn.execute("UPDATE analysis_runs SET status='cancelled' WHERE id=?", (run_id,))
        return {"run_id": run_id, "status": "cancelled"}


@mcp.tool()
def review_packet_get(run_id: str) -> dict:
    """Assemble the deterministic review packet for a queued/running run (spec §16.3):
    pinned-revision prose, declared metadata and profiles, entities on stage, prior-story
    memory, linked claims with verification state, open findings and intentional
    exceptions on the target, the seam, and the required structured-output schema.
    Marks the run 'running'. Any role (this is what the Sol runner calls)."""
    with _db() as conn:
        run = _row(conn.execute("SELECT * FROM analysis_runs WHERE id=?", (run_id,)).fetchone())
        if run is None:
            raise ValueError(f"analysis run {run_id} not found")
        if run["status"] == "cancelled":
            raise ValueError("run is cancelled")
        conn.execute("UPDATE analysis_runs SET status='running' WHERE id=?", (run_id,))
        tt, tid = run["target_type"], run["target_id"]
        packet: dict = {"run": {k: run[k] for k in ("id", "analysis_type", "target_type",
                                                    "target_id", "target_rev", "model",
                                                    "reasoning_effort", "advisory")}}
        packet["target_prose"] = _rev_content(conn, tt, tid, run["target_rev"])
        packet["target_meta"] = {r["key"]: r["value"] for r in conn.execute(
            "SELECT key, value FROM node_meta WHERE target_type=? AND target_id=?", (tt, tid))}
        pinned = json.loads(run["pinned_json"] or "{}")
        packet["profiles"] = {}
        for key, ref in pinned.items():
            prof = conn.execute("SELECT name, summary, content_md FROM entities WHERE id=?",
                                (ref["id"],)).fetchone()
            if prof is not None:
                packet["profiles"][key] = dict(prof)
        # claims linked to the target, with latest verdicts
        claims = []
        for l in conn.execute("SELECT from_id, to_id, rel_type FROM links "
                              "WHERE (from_id=? OR to_id=?) AND rel_type='asserts'", (tid, tid)):
            cid = l["to_id"] if l["from_id"] == tid else l["from_id"]
            ent = conn.execute("SELECT name, content_md FROM entities WHERE id=?",
                               (cid,)).fetchone()
            if ent is None:
                continue
            v = conn.execute("SELECT verdict, confidence FROM verifications WHERE claim_id=? "
                             "ORDER BY created_at DESC LIMIT 1", (cid,)).fetchone()
            claims.append({"id": cid, "claim": ent["content_md"] or ent["name"],
                           "verdict": v["verdict"] if v else "unchecked"})
        packet["claims"] = claims
        packet["open_findings"] = _rows(conn.execute(
            "SELECT severity, category, evidence_quote, explanation, status FROM findings "
            "WHERE target_type=? AND target_id=? AND status IN ('open','accepted','deferred')",
            (tt, tid)))
        packet["intentional_exceptions"] = _rows(conn.execute(
            "SELECT severity, category, evidence_quote, explanation, status_note FROM findings "
            "WHERE target_type=? AND target_id=? AND status='intentional'", (tt, tid)))
        packet["seam"] = seam_get(run["project_id"]) if run["advisory"] == 0 else {}
        packet["output_schema"] = SOL_OUTPUT_SCHEMA
    # context bundle opens its own connection — call outside the write transaction
    if run["advisory"] == 0 or run["analysis_type"] == "second_opinion":
        try:
            packet["context"] = (context_bundle(scene_id=tid) if tt == "scene"
                                 else context_bundle(chapter_id=tid) if tt == "chapter"
                                 else {})
        except ValueError:
            packet["context"] = {}
        if run["analysis_type"] == "second_opinion":
            # spec decision #3: outside models never see private/silhouette canon
            ctx = packet.get("context") or {}
            ctx["entities_on_stage"] = [
                e for e in ctx.get("entities_on_stage", [])
                if not str(e.get("content_md", "")).startswith("[withheld")]
            with _db() as conn2:
                hidden = {r["id"] for r in conn2.execute(
                    "SELECT id FROM entities WHERE project_id=? AND visibility IN "
                    "('private','silhouette')", (run["project_id"],))}
            ctx["entities_on_stage"] = [e for e in ctx.get("entities_on_stage", [])
                                        if e.get("id") not in hidden]
            packet["context"] = ctx
    else:
        packet["context"] = {}  # cold_read: zero context by definition
    return packet


@mcp.tool()
def analysis_run_complete(run_id: str, verdict: str, observed_summary: str,
                          intent_summary: str = "", scores: dict | None = None,
                          findings: list[dict] | None = None,
                          strengths: list[dict] | None = None,
                          limitations: list | None = None) -> dict:
    """Post a completed analysis. EVERY finding and strength must quote the pinned
    revision EXACTLY — evidence that does not appear verbatim in that revision's
    content is rejected (spec §9.6) and reported back. Valid findings open in the
    review queue. Any role (attributed)."""
    findings = findings or []
    strengths = strengths or []
    with _db() as conn:
        run = _row(conn.execute("SELECT * FROM analysis_runs WHERE id=?", (run_id,)).fetchone())
        if run is None:
            raise ValueError(f"analysis run {run_id} not found")
        if run["status"] not in ("queued", "running"):
            raise ValueError(f"run is already {run['status']}")
        content = _rev_content(conn, run["target_type"], run["target_id"], run["target_rev"])
        accepted_f, rejected = [], []
        for f in findings:
            quote = f.get("evidence_quote", "")
            sev = f.get("severity", "")
            if sev not in SEVERITIES:
                rejected.append({"reason": f"bad severity {sev!r}", **f})
                continue
            if not quote or quote not in content:
                rejected.append({"reason": "evidence_quote not found verbatim in pinned "
                                           "revision", **f})
                continue
            fid = _id()
            conn.execute(
                "INSERT INTO findings (id, run_id, project_id, target_type, target_id, "
                "target_rev, severity, category, confidence, evidence_quote, location, "
                "explanation, affected_entity_ids, smallest_intervention, created_at) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (fid, run_id, run["project_id"], run["target_type"], run["target_id"],
                 run["target_rev"], sev, f.get("category", ""), f.get("confidence", 0),
                 quote, f.get("location", ""), f.get("explanation", ""),
                 json.dumps(f.get("affected_entity_ids", [])),
                 f.get("smallest_intervention", ""), _now()))
            accepted_f.append(fid)
        accepted_s = []
        for s in strengths:
            quote = s.get("evidence_quote", "")
            if not quote or quote not in content:
                rejected.append({"reason": "strength evidence_quote not found verbatim "
                                           "in pinned revision", **s})
                continue
            sid = _id()
            conn.execute(
                "INSERT INTO strengths (id, run_id, project_id, target_type, target_id, "
                "target_rev, evidence_quote, location, explanation, created_at) "
                "VALUES (?,?,?,?,?,?,?,?,?,?)",
                (sid, run_id, run["project_id"], run["target_type"], run["target_id"],
                 run["target_rev"], quote, s.get("location", ""),
                 s.get("explanation", ""), _now()))
            accepted_s.append(sid)
        conn.execute(
            "UPDATE analysis_runs SET status='complete', verdict=?, intent_summary=?, "
            "observed_summary=?, scores_json=?, limitations_json=?, completed_by=?, "
            "completed_at=? WHERE id=?",
            (verdict, intent_summary, observed_summary, json.dumps(scores or {}),
             json.dumps(limitations or []), _caller()["name"], _now(), run_id))
        return {"run_id": run_id, "status": "complete",
                "findings_accepted": len(accepted_f), "strengths_accepted": len(accepted_s),
                "rejected": rejected}


@mcp.tool()
def annotations_get(target_type: str, target_id: str) -> dict:
    """Inline-annotation data for a reading view: live findings (open/accepted/
    deferred) and strengths-to-protect on this chapter/scene, each carrying its
    verbatim evidence quote so prose can be highlighted in place. Resolved and
    incorrect findings drop out automatically. Any role."""
    with _db() as conn:
        _get_target(conn, target_type, target_id)
        finds = _rows(conn.execute(
            "SELECT id, severity, category, evidence_quote, explanation, "
            "smallest_intervention, status, target_rev FROM findings "
            "WHERE target_type=? AND target_id=? AND status IN ('open','accepted','deferred') "
            "ORDER BY created_at", (target_type, target_id)))
        strengths, seen = [], set()
        for r in conn.execute(
                "SELECT evidence_quote, explanation, target_rev FROM strengths "
                "WHERE target_type=? AND target_id=? ORDER BY created_at DESC",
                (target_type, target_id)):
            if r["evidence_quote"] not in seen:
                seen.add(r["evidence_quote"])
                strengths.append(dict(r))
        return {"findings": finds, "strengths": strengths}


@mcp.tool()
def finding_list(project_id: str, status: str = "open",
                 target_id: str | None = None) -> list[dict]:
    """List findings (status: open | accepted | resolved | intentional | deferred |
    incorrect | all), each with a computed 'stale' flag when the target has moved past
    the finding's revision."""
    with _db() as conn:
        q = "SELECT * FROM findings WHERE project_id=?"
        args: list = [project_id]
        if status != "all":
            q += " AND status=?"
            args.append(status)
        if target_id is not None:
            q += " AND target_id=?"
            args.append(target_id)
        out = _rows(conn.execute(q + " ORDER BY created_at DESC LIMIT 200", args))
        for f in out:
            cur = conn.execute(f"SELECT rev FROM {TARGET_TABLES[f['target_type']]} WHERE id=?",
                               (f["target_id"],)).fetchone()
            f["stale"] = bool(cur) and cur["rev"] > f["target_rev"]
        return out


@mcp.tool()
def finding_update_status(finding_id: str, status: str, note: str = "") -> dict:
    """Move a finding through the workflow. resolved/accepted/deferred need author or
    owner; 'intentional' and 'incorrect' are rulings and need the OWNER (spec §3.13 —
    editors dissent via rebuttal_create instead)."""
    if status not in FINDING_STATUSES:
        raise ValueError(f"status must be one of {sorted(FINDING_STATUSES)}")
    if status in OWNER_FINDING_STATUSES:
        _require_owner()
    else:
        _require_author()
    with _db() as conn:
        if conn.execute("SELECT 1 FROM findings WHERE id=?", (finding_id,)).fetchone() is None:
            raise ValueError(f"finding {finding_id} not found")
        conn.execute("UPDATE findings SET status=?, status_note=?, updated_by=?, updated_at=? "
                     "WHERE id=?", (status, note, _caller()["name"], _now(), finding_id))
        return _row(conn.execute("SELECT * FROM findings WHERE id=?", (finding_id,)).fetchone())


@mcp.tool()
def grade_history(target_type: str, target_id: str) -> list[dict]:
    """Score trajectory for a target across completed analysis runs, oldest first —
    comparable because rubric/profile revisions are pinned per run (spec §9.7)."""
    with _db() as conn:
        _get_target(conn, target_type, target_id)
        out = []
        for r in conn.execute(
                "SELECT id, analysis_type, target_rev, verdict, scores_json, model, "
                "completed_at FROM analysis_runs WHERE target_type=? AND target_id=? "
                "AND status='complete' ORDER BY completed_at",
                (target_type, target_id)):
            out.append({"run_id": r["id"], "analysis_type": r["analysis_type"],
                        "target_rev": r["target_rev"], "verdict": r["verdict"],
                        "model": r["model"], "completed_at": r["completed_at"],
                        "scores": json.loads(r["scores_json"] or "{}")})
        return out


@mcp.tool()
def narrative_debt_list(project_id: str) -> dict:
    """The promises ledger (spec §12): open threads, unverified or false pre-seam
    claims, open blocking/major findings, watch items, stale analyses, and aging
    pending proposals. Computed live; intentional dormancy (thread status) respected."""
    with _db() as conn:
        debt: dict = {}
        threads = []
        for r in conn.execute("SELECT id, name, summary FROM entities WHERE project_id=? "
                              "AND kind='thread' AND deleted=0", (project_id,)):
            st = conn.execute("SELECT value FROM node_meta WHERE target_type='entity' "
                              "AND target_id=? AND key='status'", (r["id"],)).fetchone()
            status = st["value"] if st else "planted"
            if status not in ("paid_off", "abandoned"):
                threads.append({"name": r["name"], "status": status, "summary": r["summary"]})
        debt["open_threads"] = threads
        bad_claims = []
        for r in conn.execute("SELECT id, name FROM entities WHERE project_id=? "
                              "AND kind='research' AND deleted=0", (project_id,)):
            v = conn.execute("SELECT verdict FROM verifications WHERE claim_id=? "
                             "ORDER BY created_at DESC LIMIT 1", (r["id"],)).fetchone()
            side = conn.execute("SELECT value FROM node_meta WHERE target_type='entity' "
                                "AND target_id=? AND key='seam_side'", (r["id"],)).fetchone()
            verdict = v["verdict"] if v else "unchecked"
            if verdict in ("false", "outdated") or (
                    verdict == "unchecked" and side is not None and side["value"] == "pre"):
                bad_claims.append({"id": r["id"], "claim": r["name"], "verdict": verdict})
        debt["claims_needing_attention"] = bad_claims
        debt["open_blocking_major"] = _rows(conn.execute(
            "SELECT id, severity, category, target_id, explanation FROM findings "
            "WHERE project_id=? AND status='open' AND severity IN ('blocking','major') "
            "ORDER BY severity", (project_id,)))
        debt["watch_items"] = _rows(conn.execute(
            "SELECT id, category, explanation FROM findings WHERE project_id=? "
            "AND status='open' AND severity='watch'", (project_id,)))
        stale_runs = 0
        for r in conn.execute("SELECT target_type, target_id, target_rev FROM analysis_runs "
                              "WHERE project_id=? AND status='complete'", (project_id,)):
            cur = conn.execute(f"SELECT rev FROM {TARGET_TABLES[r['target_type']]} WHERE id=?",
                               (r["target_id"],)).fetchone()
            if cur is not None and cur["rev"] > r["target_rev"]:
                stale_runs += 1
        debt["stale_analyses"] = stale_runs
        debt["pending_proposals"] = conn.execute(
            "SELECT COUNT(*) n FROM proposals WHERE status='pending' AND target_id IN "
            "(SELECT id FROM chapters WHERE project_id=? UNION SELECT id FROM scenes "
            "WHERE project_id=? UNION SELECT id FROM entities WHERE project_id=?)",
            (project_id, project_id, project_id)).fetchone()["n"]
        return debt


@mcp.tool()
def project_dashboard_get(project_id: str) -> dict:
    """The room's dashboard, in spec §15.3 priority order: what changed, what's working
    (protect), blocking/major findings, decisions awaiting the owner, reviews waiting,
    stale analyses, and the debt summary. Never leads with one synthetic score."""
    with _db() as conn:
        if conn.execute("SELECT 1 FROM projects WHERE id=? AND deleted=0",
                        (project_id,)).fetchone() is None:
            raise ValueError(f"project {project_id} not found")
        dash: dict = {}
        changed = []
        for r in conn.execute(
                "SELECT r.target_type, r.target_id, r.rev, r.note, r.created_by, r.created_at "
                "FROM revisions r ORDER BY r.created_at DESC LIMIT 30"):
            tbl = TARGET_TABLES[r["target_type"]]
            owner_row = conn.execute(
                f"SELECT project_id, COALESCE(title, '') AS t FROM {tbl} WHERE id=?"
                if r["target_type"] != "entity" else
                f"SELECT project_id, name AS t FROM {tbl} WHERE id=?",
                (r["target_id"],)).fetchone()
            if owner_row is None or owner_row["project_id"] != project_id:
                continue
            changed.append({"what": owner_row["t"], "type": r["target_type"], "rev": r["rev"],
                            "by": r["created_by"], "at": r["created_at"], "note": r["note"]})
            if len(changed) >= 10:
                break
        dash["what_changed"] = changed
        dash["protect"] = _rows(conn.execute(
            "SELECT evidence_quote, explanation, target_id FROM strengths WHERE project_id=? "
            "ORDER BY created_at DESC LIMIT 10", (project_id,)))
        dash["blocking_major_findings"] = _rows(conn.execute(
            "SELECT id, severity, category, target_id, evidence_quote, smallest_intervention "
            "FROM findings WHERE project_id=? AND status='open' "
            "AND severity IN ('blocking','major')", (project_id,)))
        awaiting = {"pending_proposals": [], "open_blocking": len(dash["blocking_major_findings"])}
        for p in conn.execute(
                "SELECT id, target_type, target_id, created_by, created_at FROM proposals "
                "WHERE status='pending' ORDER BY created_at"):
            row = conn.execute(
                f"SELECT project_id FROM {TARGET_TABLES[p['target_type']]} WHERE id=?",
                (p["target_id"],)).fetchone()
            if row is not None and row["project_id"] == project_id:
                awaiting["pending_proposals"].append(dict(p))
        dash["awaiting_owner"] = awaiting
        dash["reviews_waiting"] = _rows(conn.execute(
            "SELECT id, analysis_type, target_id, status, created_at FROM analysis_runs "
            "WHERE project_id=? AND status IN ('queued','running') ORDER BY created_at",
            (project_id,)))
        # Latest completed gate PER MODEL (dual-review governance: Sol outside,
        # Luna inside; disagreements are the point and both belong on the Desk).
        gates, seen_models = [], set()
        for latest in conn.execute(
                "SELECT id, analysis_type, target_type, target_id, target_rev, verdict, "
                "model, completed_by, scores_json, completed_at FROM analysis_runs "
                "WHERE project_id=? AND status='complete' AND advisory=0 "
                "ORDER BY completed_at DESC", (project_id,)):
            if latest["model"] in seen_models:
                continue
            seen_models.add(latest["model"])
            cur = conn.execute(
                f"SELECT rev FROM {TARGET_TABLES[latest['target_type']]} WHERE id=?",
                (latest["target_id"],)).fetchone()
            gates.append({
                "run_id": latest["id"], "analysis_type": latest["analysis_type"],
                "verdict": latest["verdict"], "target_rev": latest["target_rev"],
                "model": latest["model"], "completed_by": latest["completed_by"],
                "completed_at": latest["completed_at"],
                "scores": json.loads(latest["scores_json"] or "{}"),
                "stale": bool(cur) and cur["rev"] > latest["target_rev"]})
        dash["latest_gates"] = gates
        dash["latest_gate"] = gates[0] if gates else None
        dash["debt"] = narrative_debt_list(project_id)
        dash["seam"] = seam_get(project_id)
        return dash


# ---------------------------------------------------------------- appearances & timeline

@mcp.tool()
def entity_appearances(entity_id: str) -> list[dict]:
    """Everywhere an entity's name/aliases appear in prose, in manuscript order —
    the 'when did we last see her' query. Maintained automatically on content writes."""
    with _db() as conn:
        ent = _get_target(conn, "entity", entity_id)
        out = _rows(conn.execute(
            "SELECT m.target_type, m.target_id, m.count, "
            "  COALESCE(c.title, s.title, e.name) AS title, "
            "  COALESCE(c.sort_order, cs.sort_order, 0) AS chapter_order, "
            "  COALESCE(s.sort_order, 0) AS scene_order "
            "FROM mentions m "
            "LEFT JOIN chapters c  ON m.target_type='chapter' AND c.id = m.target_id "
            "LEFT JOIN scenes s    ON m.target_type='scene'   AND s.id = m.target_id "
            "LEFT JOIN chapters cs ON s.chapter_id = cs.id "
            "LEFT JOIN entities e  ON m.target_type='entity'  AND e.id = m.target_id "
            "WHERE m.entity_id=? "
            "ORDER BY m.target_type='entity', chapter_order, scene_order",
            (entity_id,)))
        return [{"entity": ent["name"], **r} for r in out]


@mcp.tool()
def mentions_rebuild(project_id: str) -> dict:
    """Rescan every chapter/scene/entity in a project for entity mentions. Run after
    renaming an entity or adding aliases (meta key 'aliases', comma-separated).
    Author role required."""
    _require_author()
    n = 0
    with _db() as conn:
        for tt, table in TARGET_TABLES.items():
            for r in conn.execute(f"SELECT id FROM {table} WHERE project_id=? AND deleted=0",
                                  (project_id,)):
                _refresh_mentions(conn, tt, r["id"])
                n += 1
        return {"rescanned_nodes": n}


@mcp.tool()
def timeline_list(project_id: str) -> dict:
    """Story-time view: event entities ordered by their 'story_date' meta key
    (ISO-style strings sort correctly: '1042-03-01' or '2026-07-18T14:00'). Events
    without a story_date are listed separately. Each event includes its links so
    participants/locations are visible. Set dates via meta_set(entity, id,
    'story_date', ...). Narrative order lives in chapter/scene sort_order —
    the two orderings are deliberately independent."""
    with _db() as conn:
        events = _rows(conn.execute(
            "SELECT e.id, e.name, e.summary, m.value AS story_date "
            "FROM entities e "
            "LEFT JOIN node_meta m ON m.target_type='entity' AND m.target_id=e.id "
            "  AND m.key='story_date' "
            "WHERE e.project_id=? AND e.kind='event' AND e.deleted=0",
            (project_id,)))
        for ev in events:
            ev["links"] = _rows(conn.execute(
                "SELECT from_id, to_id, rel_type, attrs FROM links "
                "WHERE from_id=? OR to_id=?", (ev["id"], ev["id"])))
        dated = sorted([e for e in events if e["story_date"]], key=lambda e: e["story_date"])
        undated = [e for e in events if not e["story_date"]]
        return {"dated": dated, "undated": undated}


# ---------------------------------------------------------------- structure templates

STORY_TEMPLATES: dict[str, list[tuple[str, str]]] = {
    "three_act": [
        ("Setup", "Establish protagonist, world, and the want that drives them."),
        ("Inciting Incident", "The event that disrupts the status quo and demands response."),
        ("First Plot Point", "Protagonist commits; the door back closes."),
        ("Rising Complications", "Escalating obstacles; stakes compound."),
        ("Midpoint", "A reversal or revelation that reframes the goal."),
        ("Crisis", "The lowest point; the want and the need collide."),
        ("Climax", "Final confrontation; the protagonist's choice decides it."),
        ("Resolution", "The new equilibrium; show what changed."),
    ],
    "save_the_cat": [
        ("Opening Image", "A visual snapshot of the before-world."),
        ("Theme Stated", "Someone states the lesson the hero will resist."),
        ("Setup", "Hero's flaws and world; everything about to change."),
        ("Catalyst", "The life-changing event."),
        ("Debate", "Hero hesitates; can they really do this?"),
        ("Break into Two", "Hero chooses to act; enter the new world."),
        ("B Story", "The relationship that carries the theme."),
        ("Fun and Games", "The promise of the premise delivered."),
        ("Midpoint", "False victory or false defeat; stakes raise."),
        ("Bad Guys Close In", "External and internal pressure tightens."),
        ("All Is Lost", "The opposite of the opening; whiff of death."),
        ("Dark Night of the Soul", "Hero digests the loss."),
        ("Break into Three", "The theme clicks; the real solution appears."),
        ("Finale", "Hero executes the plan, transformed."),
        ("Final Image", "The after-world; opposite of the opening image."),
    ],
    "heros_journey": [
        ("Ordinary World", "The hero's normal before the adventure."),
        ("Call to Adventure", "The challenge arrives."),
        ("Refusal of the Call", "Fear and reluctance."),
        ("Meeting the Mentor", "Guidance and gifts."),
        ("Crossing the Threshold", "Commitment to the special world."),
        ("Tests, Allies, Enemies", "Learning the new world's rules."),
        ("Approach to the Inmost Cave", "Preparing for the central ordeal."),
        ("The Ordeal", "Death-and-rebirth crisis."),
        ("Reward", "Seizing the sword; the prize."),
        ("The Road Back", "Recommitment to return; chase."),
        ("Resurrection", "Final test; purified by sacrifice."),
        ("Return with the Elixir", "Home, changed, bearing the boon."),
    ],
    "story_circle": [
        ("You", "A character in their zone of comfort."),
        ("Need", "They want something."),
        ("Go", "They enter an unfamiliar situation."),
        ("Search", "They adapt to it."),
        ("Find", "They get what they wanted."),
        ("Take", "They pay a heavy price for it."),
        ("Return", "They go back to the familiar."),
        ("Change", "They have changed."),
    ],
}


@mcp.tool()
def template_list() -> dict:
    """List available story-structure templates and their beats."""
    return {name: [{"beat": b, "guidance": g} for b, g in beats]
            for name, beats in STORY_TEMPLATES.items()}


@mcp.tool()
def template_apply(project_id: str, template: str) -> dict:
    """Instantiate a structure template into a project: creates an arc entity as the
    spine plus one ordered note entity per beat (guidance as content), each linked
    'beat_of' the arc. Fill beats in as the story takes shape. Author role required."""
    _require_author()
    if template not in STORY_TEMPLATES:
        raise ValueError(f"template must be one of {sorted(STORY_TEMPLATES)}")
    who = _caller()["name"]
    with _db() as conn:
        if conn.execute("SELECT 1 FROM projects WHERE id=? AND deleted=0",
                        (project_id,)).fetchone() is None:
            raise ValueError(f"project {project_id} not found")
        arc_id = _id()
        conn.execute(
            "INSERT INTO entities (id, project_id, kind, name, summary, content_md, sort_order, "
            "created_by, created_at, updated_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (arc_id, project_id, "arc", f"Structure: {template}",
             f"{template} beat spine (template-generated)", "", 0, who, _now(), _now()),
        )
        _record_initial_revision(conn, "entity", arc_id, "", who)
        beat_ids = []
        for i, (beat, guidance) in enumerate(STORY_TEMPLATES[template], start=1):
            bid = _id()
            conn.execute(
                "INSERT INTO entities (id, project_id, kind, name, summary, content_md, sort_order, "
                "created_by, created_at, updated_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
                (bid, project_id, "note", f"Beat {i:02d}: {beat}", guidance, "", i,
                 who, _now(), _now()),
            )
            _record_initial_revision(conn, "entity", bid, "", who)
            conn.execute(
                "INSERT INTO links (id, project_id, from_id, to_id, rel_type, note, attrs, "
                "created_by, created_at) VALUES (?,?,?,?,?,?,?,?,?)",
                (_id(), project_id, bid, arc_id, "beat_of", "", "", who, _now()),
            )
            beat_ids.append(bid)
        return {"template": template, "arc_id": arc_id, "beat_ids": beat_ids}


# ---------------------------------------------------------------- stats

def _words(text: str) -> int:
    return len(text.split()) if text else 0


@mcp.tool()
def project_stats(project_id: str) -> dict:
    """Manuscript analytics: word counts (scene-aware), chapter/scene status rollups,
    POV distribution, top-mentioned entities, open comments/proposals, and target
    progress if the project has meta keys 'target_words' / 'deadline' (YYYY-MM-DD)."""
    with _db() as conn:
        if conn.execute("SELECT 1 FROM projects WHERE id=? AND deleted=0",
                        (project_id,)).fetchone() is None:
            raise ValueError(f"project {project_id} not found")
        chapters = []
        total_words = 0
        for ch in conn.execute(
                "SELECT id, title, status, content_md FROM chapters "
                "WHERE project_id=? AND deleted=0 ORDER BY sort_order, created_at",
                (project_id,)):
            scene_words = 0
            n_scenes = 0
            for sc in conn.execute(
                    "SELECT content_md FROM scenes WHERE chapter_id=? AND deleted=0",
                    (ch["id"],)):
                scene_words += _words(sc["content_md"])
                n_scenes += 1
            words = scene_words if scene_words else _words(ch["content_md"])
            total_words += words
            chapters.append({"id": ch["id"], "title": ch["title"], "status": ch["status"],
                             "words": words, "scenes": n_scenes})
        scene_status = {r["status"]: r["n"] for r in conn.execute(
            "SELECT status, COUNT(*) n FROM scenes WHERE project_id=? AND deleted=0 "
            "GROUP BY status", (project_id,))}
        pov = _rows(conn.execute(
            "SELECT e.name AS pov, COUNT(*) n FROM scenes s "
            "JOIN entities e ON e.id = s.pov_entity_id "
            "WHERE s.project_id=? AND s.deleted=0 GROUP BY e.name ORDER BY n DESC",
            (project_id,)))
        top_mentions = _rows(conn.execute(
            "SELECT e.name, e.kind, SUM(m.count) AS total FROM mentions m "
            "JOIN entities e ON e.id = m.entity_id "
            "WHERE m.project_id=? GROUP BY m.entity_id ORDER BY total DESC LIMIT 10",
            (project_id,)))
        open_comments = 0
        pending_proposals = 0
        for tt, table in TARGET_TABLES.items():
            open_comments += conn.execute(
                f"SELECT COUNT(*) n FROM comments WHERE target_type=? AND status='open' "
                f"AND target_id IN (SELECT id FROM {table} WHERE project_id=?)",
                (tt, project_id)).fetchone()["n"]
            pending_proposals += conn.execute(
                f"SELECT COUNT(*) n FROM proposals WHERE target_type=? AND status='pending' "
                f"AND target_id IN (SELECT id FROM {table} WHERE project_id=?)",
                (tt, project_id)).fetchone()["n"]
        out = {"total_words": total_words, "chapters": chapters,
               "scene_status_counts": scene_status, "pov_distribution": pov,
               "top_mentions": top_mentions, "open_comments": open_comments,
               "pending_proposals": pending_proposals}
        meta = {r["key"]: r["value"] for r in conn.execute(
            "SELECT key, value FROM node_meta WHERE target_type='project' AND target_id=?",
            (project_id,))}
        if meta.get("target_words", "").isdigit():
            target = int(meta["target_words"])
            out["target"] = {"target_words": target,
                             "words_remaining": max(0, target - total_words),
                             "pct_complete": round(100 * total_words / target, 1) if target else 0}
            if meta.get("deadline"):
                try:
                    days_left = (datetime.fromisoformat(meta["deadline"]).date()
                                 - datetime.now(timezone.utc).date()).days
                    out["target"]["deadline"] = meta["deadline"]
                    out["target"]["days_left"] = days_left
                    if days_left > 0:
                        out["target"]["words_per_day_needed"] = -(-max(
                            0, target - total_words) // days_left)
                except ValueError:
                    pass
        return out


# ---------------------------------------------------------------- export

def _assemble_manuscript(conn, project_id: str) -> tuple[str, list[tuple[str, str]]]:
    """Return (title, [(chapter_title, chapter_body_md)]) in manuscript order.
    A chapter with scenes renders its scenes joined by scene-break markers;
    a scene-less chapter uses its own content."""
    proj = conn.execute("SELECT * FROM projects WHERE id=? AND deleted=0",
                        (project_id,)).fetchone()
    if proj is None:
        raise ValueError(f"project {project_id} not found")
    parts: list[tuple[str, str]] = []
    for ch in conn.execute(
            "SELECT id, title, content_md FROM chapters WHERE project_id=? AND deleted=0 "
            "ORDER BY sort_order, created_at", (project_id,)):
        scenes = [s["content_md"] for s in conn.execute(
            "SELECT content_md FROM scenes WHERE chapter_id=? AND deleted=0 "
            "ORDER BY sort_order, created_at", (ch["id"],)) if s["content_md"].strip()]
        body = "\n\n***\n\n".join(scenes) if scenes else (ch["content_md"] or "")
        parts.append((ch["title"], body))
    return proj["name"], parts


@mcp.tool()
def export_manuscript(project_id: str, format: str = "markdown") -> dict:
    """Assemble the ordered manuscript (scene-aware) into a deliverable file:
    markdown | html | docx | epub. Returns the file name and its authed download
    path (GET /export/<file> with your API key). Any role."""
    if format not in ("markdown", "html", "docx", "epub"):
        raise ValueError("format must be markdown | html | docx | epub")
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    with _db() as conn:
        title, parts = _assemble_manuscript(conn, project_id)
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", title).strip("-").lower() or "manuscript"
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    ext = {"markdown": "md", "html": "html", "docx": "docx", "epub": "epub"}[format]
    dest = EXPORT_DIR / f"{slug}-{stamp}.{ext}"

    if format == "markdown":
        md = f"# {title}\n\n" + "\n\n".join(f"## {ct}\n\n{body}" for ct, body in parts)
        dest.write_text(md, encoding="utf-8")
    elif format == "html":
        try:
            import markdown as _md
            render = lambda t: _md.markdown(t)
        except ImportError:
            render = lambda t: "".join(f"<p>{p}</p>" for p in t.split("\n\n"))
        body_html = "".join(
            f"<h2>{ct}</h2>{render(body)}" for ct, body in parts)
        dest.write_text(
            f"<!doctype html><html><head><meta charset='utf-8'><title>{title}</title>"
            f"</head><body><h1>{title}</h1>{body_html}</body></html>", encoding="utf-8")
    elif format == "docx":
        try:
            import docx  # python-docx
        except ImportError:
            raise RuntimeError("docx export needs python-docx installed on the server")
        doc = docx.Document()
        doc.add_heading(title, level=0)
        for ct, body in parts:
            doc.add_heading(ct, level=1)
            for para in body.split("\n\n"):
                if para.strip() == "***":
                    doc.add_paragraph("* * *")
                elif para.strip():
                    doc.add_paragraph(para.strip())
        doc.save(str(dest))
    else:  # epub
        try:
            from ebooklib import epub
        except ImportError:
            raise RuntimeError("epub export needs EbookLib installed on the server")
        try:
            import markdown as _md
            render = lambda t: _md.markdown(t)
        except ImportError:
            render = lambda t: "".join(f"<p>{p}</p>" for p in t.split("\n\n"))
        book = epub.EpubBook()
        book.set_identifier(project_id)
        book.set_title(title)
        book.set_language("en")
        items = []
        for i, (ct, body) in enumerate(parts, start=1):
            c = epub.EpubHtml(title=ct, file_name=f"ch{i:03d}.xhtml", lang="en")
            c.content = f"<h1>{ct}</h1>{render(body)}"
            book.add_item(c)
            items.append(c)
        book.toc = items
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ["nav"] + items
        epub.write_epub(str(dest), book)

    return {"file": dest.name, "bytes": dest.stat().st_size, "format": format,
            "download": f"/export/{dest.name}",
            "chapters": len(parts)}


# ---------------------------------------------------------------- context assembly

@mcp.tool()
def context_bundle(chapter_id: str | None = None, scene_id: str | None = None,
                   max_prior_chars: int = 8000) -> dict:
    """Assemble the drafting context for one chapter or scene in a single call:
    project style guides, the POV character sheet, every entity on stage (mentioned
    in or linked to the target, plus 'always' entities), token-light prior-story
    memory (synopses of everything before it), and open threads.

    Honors two entity meta conventions (set via meta_set):
      ai_context = always | detected (default) | never
      reveal_after_chapter = N  (content withheld while drafting chapters before N)

    This is what an AI co-writer should read before writing. Any role."""
    if (chapter_id is None) == (scene_id is None):
        raise ValueError("provide exactly one of chapter_id or scene_id")
    with _db() as conn:
        if scene_id is not None:
            target = _get_target(conn, "scene", scene_id)
            t_type, t_id = "scene", scene_id
            chapter = _get_target(conn, "chapter", target["chapter_id"])
        else:
            target = _get_target(conn, "chapter", chapter_id)
            t_type, t_id = "chapter", chapter_id
            chapter = target
        project_id = target["project_id"]

        chapters = _rows(conn.execute(
            "SELECT id, title, content_md FROM chapters WHERE project_id=? AND deleted=0 "
            "ORDER BY sort_order, created_at", (project_id,)))
        position = next((i for i, c in enumerate(chapters, start=1)
                         if c["id"] == chapter["id"]), len(chapters))

        # entity meta conventions, one query each
        ai_modes = {r["target_id"]: r["value"] for r in conn.execute(
            "SELECT target_id, value FROM node_meta WHERE target_type='entity' "
            "AND key='ai_context'")}
        reveals = {r["target_id"]: r["value"] for r in conn.execute(
            "SELECT target_id, value FROM node_meta WHERE target_type='entity' "
            "AND key='reveal_after_chapter'")}

        on_stage_ids = {r["entity_id"] for r in conn.execute(
            "SELECT entity_id FROM mentions WHERE target_type=? AND target_id=?",
            (t_type, t_id))}
        for r in conn.execute("SELECT from_id, to_id FROM links WHERE from_id=? OR to_id=?",
                              (t_id, t_id)):
            on_stage_ids.update((r["from_id"], r["to_id"]))
        on_stage_ids.discard(t_id)
        for eid, mode in ai_modes.items():
            if mode == "always":
                on_stage_ids.add(eid)

        def _entity_payload(eid: str) -> dict | None:
            if ai_modes.get(eid) == "never":
                return None
            row = conn.execute(
                "SELECT id, kind, name, summary, content_md FROM entities "
                "WHERE id=? AND project_id=? AND deleted=0", (eid, project_id)).fetchone()
            if row is None:
                return None
            ent = dict(row)
            gate = reveals.get(eid, "")
            if gate.isdigit() and int(gate) > position:
                ent["content_md"] = f"[withheld until chapter {gate}]"
            else:
                ent["meta"] = {r["key"]: r["value"] for r in conn.execute(
                    "SELECT key, value FROM node_meta WHERE target_type='entity' "
                    "AND target_id=?", (eid,))}
            return ent

        on_stage = [e for e in (_entity_payload(i) for i in sorted(on_stage_ids)) if e]

        style = _rows(conn.execute(
            "SELECT name, summary, content_md FROM entities "
            "WHERE project_id=? AND kind='style' AND deleted=0 ORDER BY sort_order",
            (project_id,)))

        pov = None
        pov_id = target["pov_entity_id"] if t_type == "scene" else None
        if pov_id and ai_modes.get(pov_id) != "never":
            pov = _entity_payload(pov_id)

        # prior-story memory: synopses of everything before the target, newest last
        prior: list[str] = []
        used = 0
        for i, c in enumerate(chapters, start=1):
            if i >= position:
                break
            syns = [s["synopsis"] for s in conn.execute(
                "SELECT synopsis FROM scenes WHERE chapter_id=? AND deleted=0 "
                "ORDER BY sort_order, created_at", (c["id"],)) if s["synopsis"].strip()]
            text = f"[{c['title']}] " + (" ".join(syns) if syns else "(no synopsis)")
            used += len(text)
            if used > max_prior_chars:
                prior.append("[earlier chapters truncated]")
                break
            prior.append(text)
        if t_type == "scene":
            before = [s["synopsis"] or s["title"] for s in conn.execute(
                "SELECT title, synopsis, sort_order, created_at FROM scenes "
                "WHERE chapter_id=? AND deleted=0 ORDER BY sort_order, created_at",
                (chapter["id"],))]
            own = _rows(conn.execute(
                "SELECT id, title, synopsis FROM scenes WHERE chapter_id=? AND deleted=0 "
                "ORDER BY sort_order, created_at", (chapter["id"],)))
            idx = next((n for n, s in enumerate(own) if s["id"] == t_id), 0)
            if idx:
                prior.append(f"[earlier in {chapter['title']}] "
                             + " ".join(x for x in before[:idx] if x))

        threads = []
        for r in conn.execute(
                "SELECT id, name, summary FROM entities "
                "WHERE project_id=? AND kind='thread' AND deleted=0", (project_id,)):
            st = conn.execute(
                "SELECT value FROM node_meta WHERE target_type='entity' AND target_id=? "
                "AND key='status'", (r["id"],)).fetchone()
            status = st["value"] if st else "planted"
            if status not in ("paid_off", "abandoned"):
                threads.append({"name": r["name"], "summary": r["summary"], "status": status})

        return {
            "target": {"type": t_type, "id": t_id, "chapter_position": position,
                       "title": target["title"], "status": target["status"],
                       "synopsis": target["synopsis"] if t_type == "scene" else "",
                       "current_content_md": target["content_md"]},
            "style": style,
            "pov_character": pov,
            "entities_on_stage": on_stage,
            "prior_story": prior,
            "open_threads": threads,
        }


@mcp.tool()
def export_list() -> list[dict]:
    """List previously exported manuscript files available at /export/<file>. Any role."""
    if not EXPORT_DIR.is_dir():
        return []
    return [{"file": p.name, "bytes": p.stat().st_size,
             "modified": datetime.fromtimestamp(p.stat().st_mtime, timezone.utc).isoformat()}
            for p in sorted(EXPORT_DIR.iterdir()) if p.is_file()]


# ---------------------------------------------------------------- backups

_backup_lock = threading.Lock()


def _do_backup(reason: str = "scheduled") -> dict:
    """Write a consistent point-in-time snapshot via VACUUM INTO and rotate old ones.
    Serialized process-wide; microsecond filenames make collisions impossible."""
    with _backup_lock:
        BACKUP_DIR.mkdir(parents=True, exist_ok=True)
        stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S-%f")
        dest = BACKUP_DIR / f"story-{stamp}.db"
        conn = _db()
        try:
            conn.execute("VACUUM INTO ?", (str(dest),))
        finally:
            conn.close()
        rotated = []
        if BACKUP_KEEP > 0:
            for old in sorted(BACKUP_DIR.glob("story-*.db"))[:-BACKUP_KEEP]:
                old.unlink()
                rotated.append(old.name)
        return {"file": dest.name, "bytes": dest.stat().st_size, "reason": reason,
                "rotated_out": rotated, "created_at": _now()}


def _latest_backup() -> Path | None:
    if not BACKUP_DIR.is_dir():
        return None
    snaps = sorted(BACKUP_DIR.glob("story-*.db"))
    return snaps[-1] if snaps else None


def _backup_loop():
    while True:
        time.sleep(BACKUP_HOURS * 3600)
        try:
            info = _do_backup()
            print(f"[story-bible] backup {info['file']} ({info['bytes']}B)")
        except Exception as e:  # the loop must survive a bad night
            print(f"[story-bible] backup FAILED: {e}", file=sys.stderr)


@mcp.tool()
def backup_now() -> dict:
    """Snapshot the database to the on-volume backups directory immediately.
    Author role required."""
    _require_author()
    return _do_backup("manual")


@mcp.tool()
def backup_list() -> list[dict]:
    """List on-volume database snapshots, oldest first. Any role."""
    if not BACKUP_DIR.is_dir():
        return []
    return [{"file": p.name, "bytes": p.stat().st_size,
             "modified": datetime.fromtimestamp(p.stat().st_mtime, timezone.utc).isoformat()}
            for p in sorted(BACKUP_DIR.glob("story-*.db"))]


# ---------------------------------------------------------------- OAuth (minimal)

import base64
import hashlib
import html as _html
import secrets as _secrets
from datetime import timedelta

PUBLIC_URL = os.environ.get("STORYBIBLE_PUBLIC_URL",
                            "https://story-bible-production.up.railway.app")


def _oauth_meta() -> bytes:
    return json.dumps({
        "issuer": PUBLIC_URL,
        "authorization_endpoint": f"{PUBLIC_URL}/oauth/authorize",
        "token_endpoint": f"{PUBLIC_URL}/oauth/token",
        "registration_endpoint": f"{PUBLIC_URL}/oauth/register",
        "response_types_supported": ["code"],
        "grant_types_supported": ["authorization_code", "refresh_token"],
        "code_challenge_methods_supported": ["S256"],
        "token_endpoint_auth_methods_supported": ["none"],
        "scopes_supported": ["storybible"],
    }).encode()


def _resource_meta() -> bytes:
    return json.dumps({
        "resource": f"{PUBLIC_URL}/mcp",
        "authorization_servers": [PUBLIC_URL],
        "bearer_methods_supported": ["header"],
    }).encode()


async def _read_body(receive) -> bytes:
    body = b""
    while True:
        msg = await receive()
        body += msg.get("body", b"")
        if not msg.get("more_body"):
            return body


def _form(qs: str) -> dict:
    out = {}
    for pair in qs.split("&"):
        if "=" in pair:
            k, v = pair.split("=", 1)
            out[k] = urllib.parse.unquote_plus(v)
    return out


AUTHORIZE_PAGE = """<!doctype html><html><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1"><title>Story Bible — Authorize</title>
<style>body{{background:#14110d;color:#e9e1d0;font-family:'Iowan Old Style',Palatino,Georgia,serif;
display:flex;align-items:center;justify-content:center;min-height:100vh;margin:0}}
.card{{width:min(420px,90vw);padding:3rem 2.4rem;border:1px solid #332c21;text-align:center;
background:#1b1712}}h1{{font-weight:400;font-size:1.6rem;letter-spacing:.3em;text-transform:uppercase}}
.rule{{width:56px;height:1px;background:#d0a04b;margin:1.1rem auto 1.4rem}}
p{{color:#9a8f7a;font-style:italic;font-size:.95rem;margin:0 0 1.4rem}}
input{{width:100%;box-sizing:border-box;padding:.7rem .9rem;background:#14110d;border:1px solid #332c21;
color:#e9e1d0;font-family:ui-monospace,Menlo,monospace;font-size:.85rem}}
button{{margin-top:1rem;width:100%;padding:.7rem;background:#d0a04b;border:none;color:#181307;
font-size:.8rem;letter-spacing:.25em;text-transform:uppercase;cursor:pointer;font-family:ui-monospace,Menlo,monospace}}
.err{{color:#c05b35;font-size:.85rem;margin-top:.9rem}}</style></head><body><div class="card">
<h1>Story Bible</h1><div class="rule"></div>
<p>A connector is asking for access. Paste an API key to grant it — the connector
inherits that key's role.</p>
<form method="POST" action="/oauth/authorize">
<input type="hidden" name="client_id" value="{client_id}">
<input type="hidden" name="redirect_uri" value="{redirect_uri}">
<input type="hidden" name="state" value="{state}">
<input type="hidden" name="code_challenge" value="{code_challenge}">
<input type="password" name="key" placeholder="API key" autocomplete="off" autofocus>
<button type="submit">Grant access</button></form>{err}</div></body></html>"""


def _render_authorize(fields: dict, err: str = "") -> bytes:
    """Consent page with EVERY interpolated field HTML-escaped (XSS-safe)."""
    e = lambda s: _html.escape(s or "", quote=True)
    return AUTHORIZE_PAGE.format(
        client_id=e(fields.get("client_id", "")),
        redirect_uri=e(fields.get("redirect_uri", "")),
        state=e(fields.get("state", "")),
        code_challenge=e(fields.get("code_challenge", "")),
        err=f"<div class='err'>{_html.escape(err)}</div>" if err else "").encode()


def _redirect_ok(conn, client_id: str, redirect_uri: str) -> bool:
    """Exact-match against the client's registered allowlist; https or loopback only.
    Checked BEFORE the consent page renders — credentials are never solicited for
    an unregistered client or destination."""
    if not client_id or not redirect_uri:
        return False
    try:
        u = urllib.parse.urlparse(redirect_uri)
    except ValueError:
        return False
    if not (u.scheme == "https" or (u.scheme == "http"
            and u.hostname in ("localhost", "127.0.0.1"))):
        return False
    row = conn.execute("SELECT redirect_uris FROM oauth_clients WHERE client_id=?",
                       (client_id,)).fetchone()
    if row is None:
        return False
    try:
        return redirect_uri in json.loads(row["redirect_uris"])
    except ValueError:
        return False


def _mint_tokens(conn, key_name: str, role: str) -> dict:
    access = "sbt_" + _secrets.token_hex(24)
    refresh = "sbr_" + _secrets.token_hex(24)
    now = datetime.now(timezone.utc)
    conn.execute("INSERT INTO oauth_tokens (token, kind, key_name, role, created_at, expires_at) "
                 "VALUES (?,?,?,?,?,?)",
                 (access, "access", key_name, role, now.isoformat(),
                  (now + timedelta(days=30)).isoformat()))
    conn.execute("INSERT INTO oauth_tokens (token, kind, key_name, role, created_at, expires_at) "
                 "VALUES (?,?,?,?,?,?)",
                 (refresh, "refresh", key_name, role, now.isoformat(),
                  (now + timedelta(days=365)).isoformat()))
    return {"access_token": access, "token_type": "Bearer",
            "expires_in": 30 * 86400, "refresh_token": refresh, "scope": "storybible"}


def _token_identity(token: str) -> dict | None:
    if not token.startswith("sbt_"):
        return None
    conn = _db()
    try:
        row = conn.execute("SELECT key_name, role, expires_at FROM oauth_tokens "
                           "WHERE token=? AND kind='access'", (token,)).fetchone()
    finally:
        conn.close()
    if row is None or row["expires_at"] < _now():
        return None
    return {"name": row["key_name"], "role": row["role"]}


import urllib.parse


async def _handle_oauth(scope, receive, send, keys) -> bool:
    """OAuth + discovery routes (all pre-auth). Returns True if handled."""
    path = scope["path"]
    method = scope.get("method", "GET")

    async def respond(status, body, ctype=b"application/json", extra=()):
        await send({"type": "http.response.start", "status": status,
                    "headers": [(b"content-type", ctype)] + list(extra)})
        await send({"type": "http.response.body", "body": body})

    if path in ("/.well-known/oauth-authorization-server",
                "/.well-known/oauth-authorization-server/mcp"):
        await respond(200, _oauth_meta()); return True
    if path in ("/.well-known/oauth-protected-resource",
                "/.well-known/oauth-protected-resource/mcp"):
        await respond(200, _resource_meta()); return True
    if path == "/oauth/register" and method == "POST":
        try:
            reg = json.loads((await _read_body(receive)).decode() or "{}")
        except ValueError:
            reg = {}
        uris = [u for u in (reg.get("redirect_uris") or []) if isinstance(u, str)]
        ok_uris = []
        for u in uris:
            p = urllib.parse.urlparse(u)
            if p.scheme == "https" or (p.scheme == "http"
                    and p.hostname in ("localhost", "127.0.0.1")):
                ok_uris.append(u)
        if not ok_uris:
            await respond(400, b'{"error": "invalid_client_metadata", '
                               b'"error_description": "redirect_uris (https or loopback) required"}')
            return True
        client_id = "sb-" + _secrets.token_hex(8)
        conn = _db()
        try:
            with conn:
                conn.execute("INSERT INTO oauth_clients (client_id, redirect_uris, created_at) "
                             "VALUES (?,?,?)", (client_id, json.dumps(ok_uris), _now()))
        finally:
            conn.close()
        await respond(201, json.dumps({
            "client_id": client_id, "redirect_uris": ok_uris,
            "token_endpoint_auth_method": "none",
            "grant_types": ["authorization_code", "refresh_token"],
            "response_types": ["code"]}).encode())
        return True
    if path == "/oauth/authorize" and method == "GET":
        q = _form(scope.get("query_string", b"").decode())
        conn = _db()
        try:
            allowed = _redirect_ok(conn, q.get("client_id", ""), q.get("redirect_uri", ""))
        finally:
            conn.close()
        if not allowed:
            await respond(400, b"Unknown client or unregistered redirect_uri. "
                               b"No credentials were requested.",
                          b"text/plain; charset=utf-8")
            return True
        await respond(200, _render_authorize(q), b"text/html; charset=utf-8"); return True
    if path == "/oauth/authorize" and method == "POST":
        f = _form((await _read_body(receive)).decode())
        conn = _db()
        try:
            allowed = _redirect_ok(conn, f.get("client_id", ""), f.get("redirect_uri", ""))
        finally:
            conn.close()
        if not allowed:
            await respond(400, b"Unknown client or unregistered redirect_uri.",
                          b"text/plain; charset=utf-8")
            return True
        ident = keys.get(f.get("key", "").strip())
        if ident is None:
            await respond(200, _render_authorize(f, "That key didn't open it."),
                          b"text/html; charset=utf-8")
            return True
        code = "sbc_" + _secrets.token_hex(24)
        now = datetime.now(timezone.utc)
        conn = _db()
        try:
            with conn:
                conn.execute("INSERT INTO oauth_codes (code, key_name, role, challenge, "
                             "redirect_uri, client_id, created_at, expires_at) "
                             "VALUES (?,?,?,?,?,?,?,?)",
                             (code, ident["name"], ident["role"],
                              f.get("code_challenge", ""), f.get("redirect_uri", ""),
                              f.get("client_id", ""), now.isoformat(),
                              (now + timedelta(minutes=10)).isoformat()))
        finally:
            conn.close()
        sep = "&" if "?" in f.get("redirect_uri", "") else "?"
        loc = f"{f.get('redirect_uri','')}{sep}code={code}&state={urllib.parse.quote(f.get('state',''))}"
        await respond(302, b"", extra=[(b"location", loc.encode())]); return True
    if path == "/oauth/token" and method == "POST":
        f = _form((await _read_body(receive)).decode())
        conn = _db()
        try:
            with conn:
                if f.get("grant_type") == "authorization_code":
                    row = conn.execute("SELECT * FROM oauth_codes WHERE code=?",
                                       (f.get("code", ""),)).fetchone()
                    if (row is None or row["used"] or row["expires_at"] < _now()
                            or row["redirect_uri"] != f.get("redirect_uri", "")
                            or (row["client_id"] or "") != f.get("client_id", "")):
                        await respond(400, b'{"error": "invalid_grant"}'); return True
                    digest = hashlib.sha256(f.get("code_verifier", "").encode()).digest()
                    if base64.urlsafe_b64encode(digest).rstrip(b"=").decode() != row["challenge"]:
                        await respond(400, b'{"error": "invalid_grant", "error_description": "PKCE"}')
                        return True
                    conn.execute("UPDATE oauth_codes SET used=1 WHERE code=?", (row["code"],))
                    out = _mint_tokens(conn, row["key_name"], row["role"])
                    await respond(200, json.dumps(out).encode()); return True
                if f.get("grant_type") == "refresh_token":
                    row = conn.execute("SELECT * FROM oauth_tokens WHERE token=? AND kind='refresh'",
                                       (f.get("refresh_token", ""),)).fetchone()
                    if row is None or row["expires_at"] < _now():
                        await respond(400, b'{"error": "invalid_grant"}'); return True
                    out = _mint_tokens(conn, row["key_name"], row["role"])
                    await respond(200, json.dumps(out).encode()); return True
        finally:
            conn.close()
        await respond(400, b'{"error": "unsupported_grant_type"}'); return True
    return False


# ---------------------------------------------------------------- ASGI app + auth

def build_app():
    """Wrap the MCP streamable-HTTP app with API-key auth. /healthz is unauthenticated."""
    inner = mcp.streamable_http_app()
    keys = _load_keys()
    if not keys:
        print("[story-bible] FATAL: STORYBIBLE_KEYS is empty — refusing to serve unauthenticated",
              file=sys.stderr)
        sys.exit(1)

    async def app(scope, receive, send):
        if scope["type"] != "http":
            return await inner(scope, receive, send)
        if scope["path"] == "/healthz":
            await send({"type": "http.response.start", "status": 200,
                        "headers": [(b"content-type", b"application/json")]})
            await send({"type": "http.response.body", "body": b'{"ok": true}'})
            return
        if scope["path"] in ("/ui", "/ui/"):
            # Read-only viewer shell. Serves no data itself — the page calls /mcp
            # with the API key the reader enters, so auth still gates everything.
            page = Path(__file__).parent / "ui.html"
            if not page.is_file():
                await send({"type": "http.response.start", "status": 404,
                            "headers": [(b"content-type", b"application/json")]})
                await send({"type": "http.response.body", "body": b'{"error": "ui not bundled"}'})
                return
            body = page.read_bytes()
            await send({"type": "http.response.start", "status": 200,
                        "headers": [(b"content-type", b"text/html; charset=utf-8"),
                                    (b"content-length", str(len(body)).encode())]})
            await send({"type": "http.response.body", "body": body})
            return
        if await _handle_oauth(scope, receive, send, keys):
            return
        headers = {k.decode().lower(): v.decode() for k, v in scope.get("headers", [])}
        key = headers.get("x-api-key", "")
        if not key and headers.get("authorization", "").lower().startswith("bearer "):
            key = headers["authorization"][7:]
        if key.startswith("sbt_"):
            # OAuth access token: full role of the key granted at the consent screen
            ident = _token_identity(key)
            if ident is None:
                await send({"type": "http.response.start", "status": 401,
                            "headers": [(b"content-type", b"application/json"),
                                        (b"www-authenticate",
                                         f'Bearer resource_metadata="{PUBLIC_URL}/.well-known/oauth-protected-resource"'.encode())]})
                await send({"type": "http.response.body",
                            "body": b'{"error": "expired or invalid token"}'})
                return
            token = CALLER.set(ident)
            try:
                await inner(scope, receive, send)
            finally:
                CALLER.reset(token)
            return
        url_key = False
        if not key:
            # Clients whose connector UIs can't set headers (ChatGPT custom connectors)
            # may carry the key in the URL: /mcp?key=... Access logging is disabled in
            # __main__ so credentials never land in our logs (spec §18). Upstream edge
            # logs are outside our control, so this transport is HARD-CAPPED below:
            # a URL-carried key never exercises more than editor (propose/comment)
            # authority, whatever role the key itself holds.
            qs = scope.get("query_string", b"").decode()
            for pair in qs.split("&"):
                if pair.startswith("key="):
                    key = pair[4:]
                    url_key = True
                    break
        ident = keys.get(key)
        if ident is not None and url_key and ident["role"] != "editor":
            ident = {"name": ident["name"], "role": "editor"}
        if ident is None:
            await send({"type": "http.response.start", "status": 401,
                        "headers": [(b"content-type", b"application/json"),
                                    (b"www-authenticate",
                                     f'Bearer resource_metadata="{PUBLIC_URL}/.well-known/oauth-protected-resource"'.encode())]})
            await send({"type": "http.response.body",
                        "body": b'{"error": "missing or invalid API key"}'})
            return
        if scope["path"] == "/backup/latest" and scope.get("method", "GET") == "GET":
            # Offsite-pull route: streams the newest snapshot (author keys only).
            if ident["role"] != "author":
                await send({"type": "http.response.start", "status": 403,
                            "headers": [(b"content-type", b"application/json")]})
                await send({"type": "http.response.body",
                            "body": b'{"error": "author key required for backup download"}'})
                return
            if b"fresh=1" in scope.get("query_string", b"") or _latest_backup() is None:
                _do_backup("pull")
            snap = _latest_backup()
            await send({"type": "http.response.start", "status": 200,
                        "headers": [(b"content-type", b"application/octet-stream"),
                                    (b"content-disposition",
                                     f'attachment; filename="{snap.name}"'.encode()),
                                    (b"content-length", str(snap.stat().st_size).encode())]})
            with snap.open("rb") as f:
                while chunk := f.read(1 << 20):
                    await send({"type": "http.response.body", "body": chunk, "more_body": True})
            await send({"type": "http.response.body", "body": b""})
            return
        if scope["path"].startswith("/export/") and scope.get("method", "GET") == "GET":
            # Download route for export_manuscript output (any authenticated key).
            name = scope["path"][len("/export/"):]
            f = EXPORT_DIR / name
            if ("/" in name or name.startswith(".") or not f.is_file()
                    or f.resolve().parent != EXPORT_DIR.resolve()):
                await send({"type": "http.response.start", "status": 404,
                            "headers": [(b"content-type", b"application/json")]})
                await send({"type": "http.response.body",
                            "body": b'{"error": "no such export"}'})
                return
            mime = {".md": "text/markdown", ".html": "text/html",
                    ".docx": ("application/vnd.openxmlformats-officedocument"
                              ".wordprocessingml.document"),
                    ".epub": "application/epub+zip"}.get(f.suffix, "application/octet-stream")
            await send({"type": "http.response.start", "status": 200,
                        "headers": [(b"content-type", mime.encode()),
                                    (b"content-disposition",
                                     f'attachment; filename="{f.name}"'.encode()),
                                    (b"content-length", str(f.stat().st_size).encode())]})
            with f.open("rb") as fh:
                while chunk := fh.read(1 << 20):
                    await send({"type": "http.response.body", "body": chunk, "more_body": True})
            await send({"type": "http.response.body", "body": b""})
            return
        token = CALLER.set(ident)
        try:
            await inner(scope, receive, send)
        finally:
            CALLER.reset(token)

    return app


if __name__ == "__main__":
    import uvicorn

    _init_db()
    if BACKUP_HOURS > 0:
        try:
            info = _do_backup("boot")
            print(f"[story-bible] boot backup {info['file']} ({info['bytes']}B)")
        except Exception as e:
            print(f"[story-bible] boot backup failed: {e}", file=sys.stderr)
        threading.Thread(target=_backup_loop, daemon=True, name="backup-loop").start()
    print(f"[story-bible] db={DB_PATH} port={PORT} keys={len(_load_keys())}")
    # access_log off: URLs may carry API keys for header-less clients (spec §18 —
    # credentials never in logs)
    uvicorn.run(build_app(), host="0.0.0.0", port=PORT, access_log=False)
